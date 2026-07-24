from flask import Flask, request, jsonify
from flask_cors import CORS
from flask_sqlalchemy import SQLAlchemy
import PyPDF2
import docx
import nltk
from nltk.corpus import stopwords
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from google import genai
import os
import re
import json
from datetime import datetime
import secrets
from werkzeug.security import generate_password_hash, check_password_hash
from collections import Counter

# Download NLTK data
try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    nltk.download('punkt')
try:
    nltk.data.find('corpora/stopwords')
except LookupError:
    nltk.download('stopwords')

# Flask App Setup
app = Flask(__name__)
CORS(app, resources={r"/*": {"origins": "*"}})
app.config['SQLALCHEMY_DATABASE_URI'] = 'mysql+pymysql://resumeapp:123456@localhost:3306/resume_analyzer'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
db = SQLAlchemy(app)

# AI Client Setup
try:
    client = genai.Client(api_key=os.getenv('GEMINI_API_KEY', 'AIzaSyD2IJ7NLDt3sphHcqUk7G6nAxzpWjdB86E'))
    AI_AVAILABLE = True
except:
    client = None
    AI_AVAILABLE = False


# ============================================================================
# DATABASE MODELS
# ============================================================================
class User(db.Model):
    __tablename__ = 'users'
    id = db.Column(db.Integer, primary_key=True, autoincrement=True)
    email = db.Column(db.String(255), unique=True, nullable=False, index=True)
    name = db.Column(db.String(255), nullable=False)
    password = db.Column(db.String(500), nullable=False)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    analyses = db.relationship('Analysis', backref='user', lazy=True, cascade='all, delete-orphan')
    sessions = db.relationship('Session', backref='user', lazy=True, cascade='all, delete-orphan')


class Session(db.Model):
    __tablename__ = 'sessions'
    id = db.Column(db.Integer, primary_key=True, autoincrement=True)
    token = db.Column(db.String(255), unique=True, nullable=False, index=True)
    user_id = db.Column(db.Integer, db.ForeignKey('users.id'), nullable=False)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)


class Analysis(db.Model):
    __tablename__ = 'analyses'
    id = db.Column(db.Integer, primary_key=True, autoincrement=True)
    analysis_id = db.Column(db.String(64), unique=True, nullable=False, index=True)
    user_id = db.Column(db.Integer, db.ForeignKey('users.id'), nullable=False)
    company_name = db.Column(db.String(255), nullable=True)
    overall_score = db.Column(db.Integer, nullable=True)
    job_match_score = db.Column(db.Float, nullable=True)
    experience_level = db.Column(db.String(50), nullable=True)
    result_json = db.Column(db.Text, nullable=True)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)


# ============================================================================
# DATABASE MANAGER
# ============================================================================
class DatabaseManager:
    @staticmethod
    def create_user(email, name, password):
        try:
            u = User(email=email.lower().strip(), name=name.strip(), password=generate_password_hash(password))
            db.session.add(u)
            db.session.commit()
            return u
        except Exception as e:
            db.session.rollback()
            raise e

    @staticmethod
    def get_user_by_email(email):
        return User.query.filter_by(email=email.lower().strip()).first()

    @staticmethod
    def create_session(user_id):
        try:
            t = secrets.token_hex(32)
            s = Session(token=t, user_id=user_id)
            db.session.add(s)
            db.session.commit()
            return t
        except Exception as e:
            db.session.rollback()
            raise e

    @staticmethod
    def get_user_by_token(token):
        s = Session.query.filter_by(token=token).first()
        return s.user if s else None

    @staticmethod
    def delete_session(token):
        try:
            s = Session.query.filter_by(token=token).first()
            if s:
                db.session.delete(s)
                db.session.commit()
        except:
            db.session.rollback()

    @staticmethod
    def create_analysis(user_id, company_name, result):
        try:
            a = Analysis(
                analysis_id=secrets.token_hex(16),
                user_id=user_id,
                company_name=company_name,
                overall_score=result.get('overallScore', 0),
                job_match_score=result.get('jobMatchScore', 0),
                experience_level=result.get('experienceLevel', 'Unknown'),
                result_json=json.dumps(result)
            )
            db.session.add(a)
            db.session.commit()
            return a
        except Exception as e:
            db.session.rollback()
            raise e


# ============================================================================
# SKILL DEFINITIONS
# ============================================================================
SOFT_SKILLS = {
    'communication skills', 'communication', 'verbal communication', 'written communication',
    'presentation skills', 'presentations', 'presenting', 'public speaking',
    'interpersonal skills', 'interpersonal', 'stakeholder engagement', 'stakeholder management',
    'client management', 'customer focus', 'cross-functional collaboration', 'collaboration',
    'teamwork', 'team player', 'leadership', 'team leadership', 'tech lead',
    'team management', 'people management', 'mentoring', 'mentorship', 'coaching',
    'decision making', 'decision-making', 'strategic thinking', 'problem solving',
    'problem-solving', 'troubleshooting', 'critical thinking', 'analytical thinking',
    'creative thinking', 'innovation mindset', 'innovative thinking', 'innovation',
    'learning agility', 'fast learner', 'quick learner', 'adaptability',
    'resilience', 'resilient', 'ownership', 'sense of ownership', 'takes ownership',
    'attention to detail', 'detail oriented', 'detail-oriented', 'time management',
    'self-motivated', 'proactive', 'initiative', 'coordination', 'organization',
    'multitasking', 'prioritization', 'flexibility', 'work ethic', 'professionalism',
    'reliability', 'accountability', 'integrity', 'conflict resolution', 'negotiation',
    'persuasion', 'emotional intelligence', 'empathy', 'active listening',
    'business acumen', 'strategic planning', 'goal setting', 'vendor management',
    'procurement', 'budget management', 'quality focus',
}

KNOWN_SKILLS = [
    'python', 'java', 'javascript', 'typescript', 'c++', 'c#', 'ruby', 'php', 'swift', 'kotlin', 'go', 'rust',
    'scala', 'perl', 'r', 'matlab', 'dart', 'lua', 'haskell', 'elixir', 'clojure', 'erlang', 'f#', 'objective-c',
    'cobol', 'fortran', 'assembly', 'groovy', 'julia', 'solidity', 'visual basic', 'vba', 'shell scripting', 'apex',
    'react', 'angular', 'vue', 'svelte', 'next.js', 'nuxt', 'gatsby', 'htmx', 'remix', 'astro',
    'node.js', 'express', 'nest.js', 'django', 'flask', 'fastapi', 'spring boot', 'laravel', 'rails', 'asp.net', '.net',
    'gin', 'fiber', 'actix', 'rocket', 'phoenix', 'koa', 'hapi', 'strapi',
    'sql', 'mysql', 'postgresql', 'mongodb', 'redis', 'oracle', 'sqlite', 'dynamodb', 'cassandra', 'elasticsearch',
    'firebase', 'supabase', 'mariadb', 'couchdb', 'neo4j', 'influxdb', 'timescaledb', 'cockroachdb', 'snowflake',
    'bigquery', 'redshift', 'memcached', 'planetscale',
    'aws', 'azure', 'gcp', 'heroku', 'vercel', 'netlify', 'digitalocean', 'cloudflare',
    'docker', 'kubernetes', 'jenkins', 'gitlab ci', 'github actions', 'terraform', 'ansible', 'prometheus', 'grafana',
    'nginx', 'apache', 'traefik', 'puppet', 'vagrant', 'argocd', 'istio', 'envoy',
    'datadog', 'new relic', 'splunk', 'elk stack', 'logstash', 'kibana', 'fluentd', 'jaeger', 'zipkin', 'opentelemetry',
    'circleci', 'travis ci', 'bamboo', 'teamcity', 'pulumi', 'rancher',
    'machine learning', 'deep learning', 'tensorflow', 'pytorch', 'keras', 'scikit-learn', 'pandas', 'numpy', 'opencv',
    'nlp', 'langchain', 'hugging face', 'transformers', 'bert', 'gpt', 'llm', 'computer vision', 'neural networks',
    'reinforcement learning', 'generative ai', 'prompt engineering', 'rag', 'mlops', 'mlflow', 'kubeflow', 'sagemaker',
    'vertex ai', 'xgboost', 'lightgbm', 'catboost', 'natural language processing', 'sentiment analysis',
    'object detection', 'image classification', 'chatbot development', 'conversational ai',
    'data analysis', 'data visualization', 'data modeling', 'data engineering', 'data warehousing', 'data mining',
    'data pipeline', 'data lake', 'etl', 'elt', 'business intelligence', 'predictive analytics', 'statistical analysis',
    'power bi', 'tableau', 'looker', 'metabase', 'superset', 'qlik', 'google data studio',
    'microsoft excel', 'google sheets', 'jupyter notebook', 'google colab', 'databricks',
    'apache spark', 'apache flink', 'apache beam', 'apache airflow', 'prefect', 'dagster', 'dbt',
    'android', 'ios', 'react native', 'flutter', 'xamarin', 'ionic', 'swiftui', 'jetpack compose', 'kotlin multiplatform',
    'html', 'css', 'sass', 'tailwind', 'bootstrap', 'material-ui', 'webpack', 'vite', 'babel', 'rollup', 'esbuild', 'parcel',
    'websocket', 'webrtc', 'web assembly', 'seo', 'web accessibility', 'styled-components', 'chakra ui', 'ant design',
    'graphql', 'rest api', 'grpc', 'oauth', 'jwt', 'swagger', 'postman', 'api gateway', 'webhook', 'soap', 'protobuf',
    'jest', 'cypress', 'selenium', 'pytest', 'unit testing', 'integration testing', 'e2e testing', 'tdd', 'bdd',
    'mocha', 'chai', 'jasmine', 'playwright', 'puppeteer', 'junit', 'testng', 'mockito',
    'jmeter', 'k6', 'gatling', 'locust', 'test automation', 'robot framework', 'appium', 'storybook',
    'git', 'github', 'gitlab', 'bitbucket', 'jira', 'confluence', 'trello', 'asana', 'figma', 'sketch', 'adobe xd',
    'microservices', 'serverless', 'event-driven architecture', 'domain-driven design', 'clean architecture',
    'mvc', 'mvvm', 'design patterns', 'solid principles', 'system design', 'distributed systems',
    'rabbitmq', 'kafka', 'celery', 'apache pulsar', 'amazon sqs', 'amazon sns',
    'cybersecurity', 'owasp', 'ssl/tls', 'sso', 'saml', 'ldap', 'active directory', 'mfa', 'iam', 'siem',
    'zero trust', 'devsecops', 'penetration testing', 'soc2', 'gdpr', 'hipaa', 'pci dss',
    'linux', 'bash', 'powershell', 'ubuntu', 'centos', 'rhel', 'debian', 'windows server',
    'blockchain', 'ethereum', 'smart contracts', 'web3', 'hardhat', 'truffle',
    'unity', 'unreal engine', 'godot', 'opengl', 'directx', 'vulkan', 'blender', 'maya',
    'sap', 'salesforce', 'oracle erp', 'microsoft dynamics', 'workday', 'servicenow',
    'agile', 'scrum', 'kanban', 'prince2', 'pmp',
    'business analysis', 'requirements gathering', 'brd', 'frd', 'srs', 'use cases', 'user stories',
    'gap analysis', 'process mapping', 'uat', 'user acceptance testing', 'root cause analysis',
    'microsoft office', 'microsoft word', 'microsoft powerpoint', 'microsoft visio', 'microsoft project', 'google workspace',
    'ci/cd', 'continuous integration', 'continuous deployment', 'infrastructure as code', 'gitops',
    'iot', 'arduino', 'raspberry pi', 'esp32', 'mqtt', 'edge computing', 'embedded systems',
    'power apps', 'power automate', 'zapier', 'retool', 'bubble', 'webflow', 'outsystems', 'mendix',
    'rpa', 'uipath', 'automation anywhere', 'blue prism',
    'wordpress', 'drupal', 'contentful', 'sentry', 'bugsnag', 'pagerduty',
    'performance optimization', 'query optimization', 'lazy loading', 'code splitting',
    'responsive design', 'progressive web apps', 'technical writing',
]

SKILL_VARIATIONS = {
    'react': ['react', 'reactjs', 'react.js'],
    'vue': ['vue', 'vuejs', 'vue.js'],
    'angular': ['angular', 'angularjs', 'angular.js'],
    'node.js': ['node.js', 'nodejs', 'node js'],
    'next.js': ['next.js', 'nextjs', 'next js'],
    'nest.js': ['nest.js', 'nestjs', 'nest js'],
    'express': ['express', 'express.js', 'expressjs'],
    'typescript': ['typescript', 'ts'],
    'javascript': ['javascript', 'js', 'es6', 'ecmascript'],
    'python': ['python', 'python3', 'python 3'],
    'c++': ['c++', 'cpp', 'cplusplus', 'c plus plus'],
    'c#': ['c#', 'csharp', 'c sharp'],
    '.net': ['.net', 'dotnet', 'asp.net', 'dot net'],
    'postgresql': ['postgresql', 'postgres', 'psql'],
    'mongodb': ['mongodb', 'mongo'],
    'kubernetes': ['kubernetes', 'k8s'],
    'aws': ['aws', 'amazon web services'],
    'gcp': ['gcp', 'google cloud', 'google cloud platform'],
    'azure': ['azure', 'microsoft azure'],
    'machine learning': ['machine learning', 'ml'],
    'deep learning': ['deep learning', 'dl'],
    'tensorflow': ['tensorflow', 'tf'],
    'scikit-learn': ['scikit-learn', 'sklearn', 'scikit learn'],
    'react native': ['react native', 'react-native', 'reactnative'],
    'spring boot': ['spring boot', 'springboot', 'spring-boot', 'spring framework', 'spring mvc', 'spring cloud', 'spring security'],
    'rest api': ['rest api', 'restful api', 'rest apis', 'restful apis', 'restful services', 'restful'],
    'ci/cd': ['ci/cd', 'cicd', 'ci cd', 'ci-cd'],
    'github actions': ['github actions', 'github action'],
    'gitlab ci': ['gitlab ci', 'gitlab-ci', 'gitlab ci/cd'],
    'power bi': ['power bi', 'powerbi', 'power-bi', 'microsoft power bi'],
    'tableau': ['tableau', 'tableau desktop', 'tableau server'],
    'microsoft excel': ['microsoft excel', 'ms excel', 'advanced excel', 'excel vba', 'excel macro', 'excel pivot'],
    'google sheets': ['google sheets', 'gsheets'],
    'data analysis': ['data analysis', 'data analytics'],
    'data visualization': ['data visualization', 'data viz'],
    'business intelligence': ['business intelligence', 'bi'],
    'predictive analytics': ['predictive analytics', 'predictive modeling'],
    'etl': ['etl', 'extract transform load'],
    'data pipeline': ['data pipeline', 'data pipelines'],
    'data warehousing': ['data warehousing', 'data warehouse', 'dwh'],
    'data lake': ['data lake', 'datalake'],
    'apache spark': ['apache spark', 'spark', 'pyspark'],
    'apache airflow': ['apache airflow', 'airflow'],
    'dbt': ['dbt', 'data build tool'],
    'snowflake': ['snowflake'],
    'bigquery': ['bigquery', 'big query', 'google bigquery'],
    'redshift': ['redshift', 'amazon redshift'],
    'databricks': ['databricks'],
    'business analysis': ['business analysis', 'ba'],
    'requirements gathering': ['requirements gathering', 'requirement gathering'],
    'brd': ['brd', 'business requirements document'],
    'frd': ['frd', 'functional requirements document'],
    'user stories': ['user stories', 'user story'],
    'use cases': ['use cases', 'use case'],
    'uat': ['uat', 'user acceptance testing'],
    'gap analysis': ['gap analysis', 'fit-gap analysis'],
    'process mapping': ['process mapping', 'process flow'],
    'root cause analysis': ['root cause analysis', 'rca'],
    'microsoft office': ['microsoft office', 'ms office', 'office 365'],
    'microsoft word': ['microsoft word', 'ms word'],
    'microsoft powerpoint': ['microsoft powerpoint', 'ms powerpoint', 'powerpoint', 'ppt'],
    'microsoft visio': ['microsoft visio', 'ms visio', 'visio'],
    'agile': ['agile', 'agile methodology', 'agile development', 'agile framework', 'agile practices', 'agile scrum'],
    'scrum': ['scrum', 'scrum master', 'scrum framework'],
    'kanban': ['kanban'],
    'cybersecurity': ['cybersecurity', 'cyber security', 'information security', 'infosec'],
    'owasp': ['owasp', 'owasp top 10'],
    'ssl/tls': ['ssl/tls', 'ssl', 'tls'],
    'sso': ['sso', 'single sign-on', 'single sign on'],
    'mfa': ['mfa', 'multi-factor authentication', '2fa', 'two factor authentication'],
    'iam': ['iam', 'identity and access management'],
    'gdpr': ['gdpr'],
    'soc2': ['soc2', 'soc 2'],
    'devsecops': ['devsecops', 'dev sec ops'],
    'elk stack': ['elk stack', 'elk', 'elastic stack'],
    'infrastructure as code': ['infrastructure as code', 'iac'],
    'gitops': ['gitops', 'git ops'],
    'argocd': ['argocd', 'argo cd', 'argo-cd'],
    'unit testing': ['unit testing', 'unit tests', 'unit test'],
    'integration testing': ['integration testing', 'integration tests'],
    'e2e testing': ['e2e testing', 'end-to-end testing', 'end to end testing', 'e2e tests'],
    'tdd': ['tdd', 'test driven development', 'test-driven development'],
    'bdd': ['bdd', 'behavior driven development', 'behaviour driven development'],
    'test automation': ['test automation', 'automated testing', 'automation testing'],
    'penetration testing': ['penetration testing', 'pen testing', 'pentest', 'ethical hacking'],
    'microservices': ['microservices', 'microservice', 'micro-services', 'micro services'],
    'serverless': ['serverless', 'faas', 'function as a service'],
    'event-driven architecture': ['event-driven architecture', 'event driven architecture', 'eda'],
    'domain-driven design': ['domain-driven design', 'domain driven design', 'ddd'],
    'system design': ['system design', 'systems design'],
    'distributed systems': ['distributed systems', 'distributed computing'],
    'design patterns': ['design patterns', 'design pattern'],
    'solid principles': ['solid principles', 'solid'],
    'clean architecture': ['clean architecture'],
    'rpa': ['rpa', 'robotic process automation'],
    'uipath': ['uipath', 'ui path'],
    'automation anywhere': ['automation anywhere'],
    'blue prism': ['blue prism', 'blueprism'],
    'power automate': ['power automate', 'microsoft power automate'],
    'iot': ['iot', 'internet of things'],
    'raspberry pi': ['raspberry pi', 'raspberrypi', 'rpi'],
    'edge computing': ['edge computing'],
    'embedded systems': ['embedded systems', 'embedded'],
    'power apps': ['power apps', 'powerapps'],
    'sap': ['sap', 'sap erp', 'sap s/4hana'],
    'salesforce': ['salesforce', 'sfdc'],
    'servicenow': ['servicenow', 'service now'],
    'blockchain': ['blockchain', 'block chain'],
    'ethereum': ['ethereum', 'eth'],
    'smart contracts': ['smart contracts', 'smart contract'],
    'web3': ['web3', 'web 3.0', 'web 3'],
    'unity': ['unity', 'unity3d', 'unity 3d', 'unity engine', 'unity game', 'unity developer'],
    'blender': ['blender', 'blender 3d', 'blender modeling', 'blender animation'],
    'maya': ['maya', 'autodesk maya', 'maya 3d', 'maya modeling'],
    'sentry': ['sentry', 'sentry.io'],
    'datadog': ['datadog', 'data dog'],
    'new relic': ['new relic', 'newrelic'],
    'performance optimization': ['performance optimization', 'performance tuning'],
    'query optimization': ['query optimization', 'sql optimization', 'query tuning'],
    'lazy loading': ['lazy loading'],
    'code splitting': ['code splitting'],
    'generative ai': ['generative ai', 'genai', 'gen ai'],
    'prompt engineering': ['prompt engineering'],
    'llm': ['llm', 'large language model', 'large language models'],
    'rag': ['rag', 'retrieval augmented generation'],
    'hugging face': ['hugging face', 'huggingface'],
    'chatbot development': ['chatbot development', 'chatbot', 'chat bot'],
    'mlops': ['mlops', 'ml ops'],
    'responsive design': ['responsive design', 'responsive web design'],
    'progressive web apps': ['progressive web app', 'progressive web apps', 'pwa'],
    'wordpress': ['wordpress', 'word press'],
    'web accessibility': ['web accessibility', 'a11y', 'accessibility'],
    'seo': ['seo', 'search engine optimization'],
    'swagger': ['swagger', 'openapi', 'open api'],
    'postman': ['postman'],
    'technical writing': ['technical writing', 'tech writing'],
}

STRICT_MATCH_SKILLS = {
    'go', 'r', 'c', 'dart', 'ruby', 'rust', 'swift', 'scala', 'perl', 'express', 'flask', 'rails',
    'ionic', 'redux', 'helm', 'kafka', 'nest.js', 'nuxt', 'svelte', 'grpc', 'oauth', 'jwt', 'vite',
    'babel', 'jest', 'cypress', 'celery', 'bash', 'gin', 'fiber', 'rocket', 'phoenix', 'koa', 'hapi',
    'sass', 'consul', 'vault', 'flux', 'nats', 'expo', 'remix', 'astro', 'apex', 'lua', 'julia',
    'agile', 'node.js', 'sentry', 'mqtt', 'sql', 'css', 'html', 'unity', 'blender', 'maya', 'puppet', 'vagrant',
}

CONTEXT_REQUIRED_SKILLS = {
    'go', 'scala', 'rust', 'swift', 'dart', 'ruby', 'perl', 'express', 'flask', 'rails',
    'phoenix', 'rocket', 'gin', 'fiber', 'koa', 'hapi', 'consul', 'vault', 'flux', 'nats', 'expo', 'remix',
    'agile', 'apex', 'lua', 'julia', 'helm', 'celery', 'babel', 'sass', 'ionic', 'redux', 'kafka', 'mqtt',
    'unity', 'blender', 'maya', 'puppet', 'vagrant',
}

FALSE_POSITIVE_PATTERNS = {
    'go': ['go-live', 'golive', 'go live', 'going', 'good', 'goal', 'goals', 'gone', 'got', 'govern', 'governance', 'google', 'goods', 'ago', 'undergo', 'ego', 'ergo', 'logo', 'mango', 'category', 'cargo', 'ongoing', 'outgoing', 'forgo'],
    'r': ['brd', 'frd', 'crp', 'uat', 'erp', 'mr', 'dr', 'sr', 'jr', 'pr', 'or', 'for', 'are', 'our', 'your', 'their', 'her', 'per', 'after', 'under', 'over', 'other', 'every', 'where', 'there', 'here', 'more', 'before', 'never', 'ever', 'during', 'from'],
    'scala': ['scalable', 'scalability', 'scaling', 'escalate', 'escalating', 'escalation', 'scale', 'scaled'],
    'rust': ['trust', 'trusted', 'trustworthy', 'robust', 'industry', 'frustrate', 'frustrated', 'entrust'],
    'swift': ['swiftly', 'swiftness'],
    'dart': ['standard', 'standards', 'darting', 'dashboard'],
    'ruby': ['rubric', 'rubrics'],
    'express': ['express interest', 'express ideas', 'expressed', 'expressing', 'expression', 'expressly', 'express their', 'express your', 'express the', 'express concern', 'express delivery'],
    'flask': ['flashback'],
    'rails': ['trail', 'trails', 'derail', 'derails', 'guardrails'],
    'helm': ['help', 'helped', 'helping', 'helpful', 'overwhelm', 'overwhelming'],
    'ionic': ['electronic', 'electronics', 'chronicle', 'chronically'],
    'sass': ['saas', 'sas'],
    'redux': ['reduce', 'reduced', 'reducing', 'reduction'],
    'perl': ['experience', 'expert', 'expertise', 'preferably', 'properly', 'superlative', 'interpersonal'],
    'jest': ['suggest', 'suggestion', 'majestic', 'majesty'],
    'celery': ['salary', 'accelery'],
    'vite': ['invite', 'invited', 'invitation', 'favorite', 'favourite'],
    'bash': ['bashing', 'abash'],
    'consul': ['consult', 'consultant', 'consulting', 'consultation'],
    'vault': ['default', 'defaults'],
    'flux': ['influx'],
    'nats': ['natural', 'naturally', 'national'],
    'expo': ['expose', 'exposed', 'exposure', 'export', 'exported'],
    'gin': ['going', 'beginning', 'engineering', 'managing', 'changing', 'emerging', 'designing', 'login', 'margin', 'origin'],
    'fiber': ['fibers'],
    'hapi': ['happy', 'happiness', 'happily'],
    'lua': ['value', 'evaluate', 'evaluation', 'valuable'],
    'css': ['access', 'accessing', 'success', 'successful', 'necessary'],
    'puppet': ['puppet show', 'puppeteer'],
    'unity': ['community', 'opportunity', 'unit'],
    'blender': ['blended', 'blending'],
    'vagrant': ['vagrancy'],
    'agile': [], 'phoenix': [], 'koa': [], 'rocket': [], 'remix': [], 'astro': [],
    'apex': [], 'julia': [], 'kafka': [], 'nuxt': [], 'svelte': [],
    'grpc': [], 'oauth': [], 'jwt': [], 'babel': [], 'cypress': [],
    'nest.js': [], 'node.js': [], 'sentry': [], 'mqtt': [], 'sql': [], 'html': [], 'maya': [],
}

TECH_CONTEXT_WORDS = [
    'programming', 'language', 'framework', 'library', 'developer', 'development', 'software', 'code', 'coding',
    'engineer', 'engineering', 'stack', 'backend', 'frontend', 'fullstack', 'full-stack', 'api', 'database', 'server',
    'cloud', 'deploy', 'deployment', 'devops', 'container', 'microservice', 'application', 'script', 'scripting',
    'testing', 'debug', 'compile', 'compiler', 'runtime', 'sdk', 'ide', 'package', 'algorithm', 'machine learning',
    'ai', 'ml', 'repository', 'version control', 'ci/cd', 'build', 'release', 'infrastructure', 'kubernetes',
    'docker', 'aws', 'azure', 'gcp', 'proficient', 'expertise', 'skilled', 'technologies', 'tools', 'platform',
    'architecture', 'experience with', 'knowledge of', 'proficiency', 'tech stack', 'technical', 'integration',
    'react', 'angular', 'vue', 'python', 'java', 'node', 'django', 'fastapi', 'spring boot',
    'terraform', 'jenkins', 'git', 'github', 'gitlab',
]

RELATED_SKILL_GROUPS = {
    'frontend': ['react', 'angular', 'vue', 'svelte', 'next.js', 'nuxt', 'gatsby', 'remix', 'astro'],
    'backend': ['express', 'django', 'flask', 'fastapi', 'spring boot', 'nest.js', 'laravel', 'rails', 'gin', 'fiber', 'phoenix', 'koa'],
    'languages': ['python', 'java', 'javascript', 'typescript', 'c++', 'c#', 'go', 'rust', 'ruby', 'kotlin', 'swift', 'php'],
    'sql_db': ['sql', 'mysql', 'postgresql', 'oracle', 'sqlite', 'mariadb', 'snowflake', 'bigquery', 'redshift'],
    'nosql_db': ['mongodb', 'redis', 'dynamodb', 'cassandra', 'firebase', 'couchdb', 'neo4j'],
    'cloud': ['aws', 'azure', 'gcp'],
    'containers': ['docker', 'kubernetes', 'rancher'],
    'ci_cd': ['jenkins', 'gitlab ci', 'github actions', 'ci/cd', 'circleci', 'travis ci', 'bamboo', 'teamcity'],
    'iac': ['terraform', 'ansible', 'puppet', 'pulumi'],
    'monitoring': ['prometheus', 'grafana', 'datadog', 'new relic', 'splunk', 'elk stack', 'jaeger', 'zipkin'],
    'css_fw': ['css', 'sass', 'tailwind', 'bootstrap', 'material-ui', 'chakra ui', 'ant design', 'styled-components'],
    'testing': ['jest', 'cypress', 'selenium', 'pytest', 'playwright', 'puppeteer', 'junit', 'testng', 'mocha'],
    'mobile': ['android', 'ios', 'react native', 'flutter', 'xamarin', 'ionic', 'swiftui', 'jetpack compose'],
    'ml': ['machine learning', 'deep learning', 'tensorflow', 'pytorch', 'keras', 'scikit-learn', 'xgboost', 'lightgbm'],
    'data': ['pandas', 'numpy', 'r', 'matlab', 'jupyter notebook', 'google colab'],
    'data_proc': ['apache spark', 'apache flink', 'apache beam', 'apache airflow', 'dbt', 'prefect', 'dagster'],
    'bi': ['power bi', 'tableau', 'looker', 'metabase', 'superset', 'qlik', 'google data studio'],
    'messaging': ['rabbitmq', 'kafka', 'celery', 'amazon sqs', 'amazon sns', 'apache pulsar'],
    'api': ['rest api', 'graphql', 'grpc', 'websocket', 'soap'],
    'vcs': ['git', 'github', 'gitlab', 'bitbucket'],
    'auth': ['oauth', 'jwt', 'sso', 'saml', 'ldap'],
    'pm': ['jira', 'confluence', 'trello', 'asana'],
    'design': ['figma', 'sketch', 'adobe xd'],
    'security': ['cybersecurity', 'owasp', 'penetration testing', 'devsecops', 'siem'],
    'methodology': ['agile', 'scrum', 'kanban'],
    'ba': ['business analysis', 'requirements gathering', 'gap analysis', 'process mapping', 'uat', 'brd', 'frd', 'user stories', 'use cases'],
    'erp': ['sap', 'salesforce', 'oracle erp', 'microsoft dynamics', 'workday', 'servicenow'],
    'office': ['microsoft office', 'microsoft excel', 'microsoft word', 'microsoft powerpoint', 'google workspace', 'google sheets'],
    'rpa': ['rpa', 'uipath', 'automation anywhere', 'blue prism', 'power automate'],
    'blockchain': ['blockchain', 'ethereum', 'solidity', 'smart contracts', 'web3'],
    'ai_adv': ['generative ai', 'prompt engineering', 'llm', 'rag', 'langchain', 'hugging face', 'chatbot development'],
    'arch': ['microservices', 'serverless', 'event-driven architecture', 'domain-driven design', 'system design', 'distributed systems', 'clean architecture', 'design patterns'],
    'spreadsheets': ['microsoft excel', 'google sheets'],
    'iot': ['iot', 'arduino', 'raspberry pi', 'mqtt', 'edge computing', 'embedded systems'],
    'lowcode': ['power apps', 'retool', 'bubble', 'webflow', 'outsystems', 'mendix'],
}

KNOWN_CERTIFICATIONS = [
    'aws certified', 'aws solutions architect', 'aws developer', 'aws sysops', 'aws cloud practitioner',
    'aws devops engineer', 'aws machine learning', 'azure certified', 'azure fundamentals', 'az-900', 'az-104',
    'az-204', 'az-400', 'google cloud certified', 'gcp professional', 'google cloud engineer',
    'certified kubernetes administrator', 'cka', 'ckad', 'cks', 'terraform certified', 'hashicorp certified',
    'pmp', 'project management professional', 'prince2', 'prince2 certified',
    'scrum master', 'csm', 'psm', 'cissp', 'cism', 'ceh', 'certified ethical hacker',
    'comptia security+', 'comptia network+', 'comptia a+', 'ccna', 'ccnp', 'cisco certified',
    'oracle certified', 'ocp', 'oca', 'microsoft certified', 'mcsa', 'mcse',
    'salesforce certified', 'salesforce administrator', 'sap certified',
    'istqb', 'istqb certified', 'itil certified', 'itil foundation',
    'rhce', 'red hat certified', 'docker certified', 'dca',
    'certified scrum developer', 'csd', 'togaf certified',
    'six sigma', 'six sigma green belt', 'six sigma black belt',
    'tensorflow developer certificate', 'tensorflow certified',
    'databricks certified', 'snowflake certified', 'mongodb certified',
    'elastic certified', 'linux foundation certified', 'lfcs', 'lfce',
    'certified information systems auditor', 'cisa', 'google analytics certified',
    'meta certified',
]


# ============================================================================
# UTILITY FUNCTIONS
# ============================================================================
def safe_str(v, d=""):
    return str(v) if v is not None else d

def safe_list(v):
    return v if isinstance(v, list) else []

def safe_dict(v):
    return v if isinstance(v, dict) else {}

def safe_int(v, d=0):
    try:
        return int(v)
    except:
        return d

def is_soft_skill(s):
    return s.lower().strip() in SOFT_SKILLS if s else False

def get_skill_resources_data(skill):
    if not skill:
        skill = "Unknown"
    dn = skill.title()
    return {
        'displayName': dn,
        'category': 'Technical Skill',
        'difficulty': 'Medium',
        'learningTime': '4-6 weeks',
        'salaryImpact': '+$10K-$20K',
        'demandLevel': 'High',
        'description': f'{dn} is a valuable technical skill.',
        'careerPaths': ['Software Developer'],
        'relatedSkills': [],
        'freeCoursePlatforms': [
            {'platform': 'freeCodeCamp', 'courseName': f'Learn {dn}', 'duration': '10-20h', 'level': 'Beginner', 'link': f'https://www.youtube.com/results?search_query=freecodecamp+{skill}', 'rating': '4.8/5'},
            {'platform': 'Udemy', 'courseName': f'{dn} Guide', 'duration': '20-40h', 'level': 'All', 'link': f'https://www.udemy.com/courses/search/?q={skill}', 'rating': '4.6/5'},
        ],
        'youtubeChannels': [
            {'channelName': 'freeCodeCamp.org', 'videoTitle': f'{dn} Course', 'duration': '4-8h', 'subscribers': '9M+', 'link': f'https://www.youtube.com/results?search_query=freecodecamp+{skill}'},
        ],
        'practiceWebsites': [{'platform': 'LeetCode', 'description': f'{dn} practice', 'difficulty': 'Medium-Hard', 'link': 'https://leetcode.com', 'type': 'Practice'}],
        'projectIdeas': [
            {'projectName': f'{dn} Starter', 'difficulty': 'Beginner', 'duration': '1 week', 'description': f'Learn {dn}', 'skills': [dn], 'whatYouLearn': f'Core {dn}'},
            {'projectName': f'{dn} Portfolio', 'difficulty': 'Intermediate', 'duration': '2 weeks', 'description': f'Showcase {dn}', 'skills': [dn], 'whatYouLearn': f'{dn} patterns'},
        ],
        'learningRoadmap': {'week1': f'{dn} fundamentals', 'week2': 'First project', 'week3': 'Intermediate', 'week4': 'Portfolio'},
        'certifications': []
    }


# ============================================================================
# RESUME ANALYZER CLASS
# ============================================================================
class ResumeAnalyzer:
    def __init__(self, ai_client=None):
        self.client = ai_client
        self.vectorizer = TfidfVectorizer(stop_words='english', max_features=1000)
        self.month_map = {
            'jan': 1, 'january': 1, 'feb': 2, 'february': 2, 'mar': 3, 'march': 3,
            'apr': 4, 'april': 4, 'may': 5, 'jun': 6, 'june': 6, 'jul': 7, 'july': 7,
            'aug': 8, 'august': 8, 'sep': 9, 'sept': 9, 'september': 9,
            'oct': 10, 'october': 10, 'nov': 11, 'november': 11, 'dec': 12, 'december': 12
        }

    # =======================================================================
    # TEXT EXTRACTION
    # =======================================================================
    def extract_pdf_text(self, file):
        try:
            reader = PyPDF2.PdfReader(file)
            return "\n".join([(p.extract_text() or "") for p in reader.pages])
        except:
            return ""

    def extract_docx_text(self, file):
        try:
            doc = docx.Document(file)
            text = "\n".join([p.text for p in doc.paragraphs])
            for t in doc.tables:
                for r in t.rows:
                    for c in r.cells:
                        text += " " + c.text
            return text
        except:
            return ""

    # =======================================================================
    # SKILL MATCHING
    # =======================================================================
    def is_false_positive(self, skill, text_lower, pos):
        for fp in FALSE_POSITIVE_PATTERNS.get(skill.lower(), []):
            if fp in text_lower:
                for m in re.finditer(re.escape(fp), text_lower):
                    if m.start() <= pos < m.start() + len(fp):
                        return True
        return False

    def skill_exists_strict(self, skill, text):
        text_lower = text.lower()
        skill_lower = skill.lower()
        variations = SKILL_VARIATIONS.get(skill_lower, [skill_lower])
        is_strict = skill_lower in STRICT_MATCH_SKILLS
        needs_ctx = skill_lower in CONTEXT_REQUIRED_SKILLS

        for var in variations:
            if len(var) <= 2:
                if self._match_single_letter(var, skill_lower, text_lower):
                    return True
                continue

            pattern = r'\b' + re.escape(var) + r'\b'
            matches = list(re.finditer(pattern, text_lower))
            if not matches:
                continue

            for match in matches:
                pos, end_pos = match.start(), match.end()
                if self.is_false_positive(skill_lower, text_lower, pos):
                    continue
                if is_strict:
                    cb = text_lower[pos - 1] if pos > 0 else ' '
                    ca = text_lower[end_pos] if end_pos < len(text_lower) else ' '
                    if cb.isalpha() or ca.isalpha():
                        continue
                if needs_ctx:
                    ctx = text_lower[max(0, pos - 80):min(len(text_lower), end_pos + 80)]
                    if any(tw in ctx for tw in TECH_CONTEXT_WORDS) or any(p in ctx for p in [',', '/', '|', '•', '·', 'skills:', 'technologies:', 'tech stack:', 'proficient in', 'experience with', 'knowledge of', 'familiar with', 'worked with', 'hands-on', 'expertise in', 'required:', 'requirements:', 'qualifications:', 'must have', 'nice to have']):
                        return True
                    continue
                return True
        return False

    def _match_single_letter(self, var, skill_name, text_lower):
        pats = {
            'r': [r'\br\s+programming', r'\br\s+language', r'\br\s+studio', r'rstudio', r'programming\s+in\s+r\b', r'language:\s*r\b', r'r,\s*(?:python|matlab|sas|spss)', r'(?:python|matlab|sas|spss),\s*r\b', r'\br\s+(?:shiny|tidyverse|ggplot|dplyr|cran)'],
            'c': [r'\bc\s+programming', r'\bc\s+language', r'programming\s+in\s+c\b', r'\bc/c\+\+', r'\bc,\s*c\+\+', r'language:\s*c\b']
        }
        for p in pats.get(skill_name, []):
            if re.search(p, text_lower):
                return True
        return False

    def extract_skills_from_text(self, text):
        found = []
        for skill in KNOWN_SKILLS:
            if is_soft_skill(skill):
                continue
            if self.skill_exists_strict(skill, text):
                if skill not in found:
                    found.append(skill)
        return found

    def match_skills(self, job_skills, resume_text):
        matched, missing = [], []
        for skill in job_skills:
            if is_soft_skill(skill):
                continue
            if self.skill_exists_strict(skill, resume_text):
                matched.append(skill)
            else:
                missing.append(skill)
        return matched, missing

    # =======================================================================
    # DETAILED EXTRACTION METHODS
    # =======================================================================
    def get_skill_match_details(self, job_skills, resume_text):
        details = []
        tl = resume_text.lower()
        for skill in job_skills:
            if is_soft_skill(skill):
                continue
            found = self.skill_exists_strict(skill, resume_text)
            context_snippet = ""
            if found:
                variations = SKILL_VARIATIONS.get(skill.lower(), [skill.lower()])
                for var in variations:
                    if len(var) <= 2:
                        continue
                    m = re.search(r'\b' + re.escape(var) + r'\b', tl)
                    if m:
                        start = max(0, m.start() - 60)
                        end = min(len(tl), m.end() + 60)
                        context_snippet = "..." + resume_text[start:end].strip().replace('\n', ' ') + "..."
                        break
            details.append({
                "skill": skill,
                "found": found,
                "status": "✅ Matched" if found else "❌ Missing",
                "contextInResume": context_snippet if found else "Not found in resume"
            })
        return details

    def get_skill_depth_details(self, matched_skills, resume_text):
        tl = resume_text.lower()
        project_words = ['built', 'developed', 'created', 'designed', 'implemented', 'architected', 'deployed',
                        'migrated', 'integrated', 'optimized', 'led', 'managed', 'delivered', 'launched',
                        'automated', 'project', 'application', 'system', 'platform', 'service']
        details = []
        for skill in matched_skills:
            variations = SKILL_VARIATIONS.get(skill.lower(), [skill.lower()])
            total_count = 0
            in_project = False
            snippets = []
            for var in variations:
                if len(var) <= 2:
                    continue
                for m in re.finditer(r'\b' + re.escape(var) + r'\b', tl):
                    total_count += 1
                    ctx_s = max(0, m.start() - 80)
                    ctx_e = min(len(tl), m.end() + 80)
                    ctx = tl[ctx_s:ctx_e]
                    if any(pw in ctx for pw in project_words):
                        in_project = True
                    if len(snippets) < 3:
                        snippets.append("..." + resume_text[ctx_s:ctx_e].strip().replace('\n', ' ') + "...")

            if in_project and total_count >= 3:
                depth = "Expert"
            elif total_count >= 3:
                depth = "Experienced"
            elif in_project:
                depth = "Practical"
            elif total_count >= 2:
                depth = "Familiar"
            else:
                depth = "Mentioned"

            score = 3 if depth == "Expert" else 2.5 if depth == "Experienced" else 2 if depth == "Practical" else 1.5 if depth == "Familiar" else 1
            details.append({
                "skill": skill,
                "mentionCount": total_count,
                "usedInProject": in_project,
                "depthLevel": depth,
                "depthScore": f"{score}/3",
                "contextSnippets": snippets
            })
        avg = sum(d['mentionCount'] for d in details) / max(len(details), 1)
        return details, round(avg, 1)

    def get_transferable_details(self, missing_skills, resume_skills):
        details = []
        for ms in missing_skills:
            ml = ms.lower()
            related = []
            group_name = ""
            for gn, gs in RELATED_SKILL_GROUPS.items():
                if ml in gs:
                    group_name = gn
                    for rs in resume_skills:
                        if rs.lower() in gs and rs.lower() != ml:
                            related.append(rs)
            details.append({
                "missingSkill": ms,
                "skillGroup": group_name,
                "relatedSkillsYouHave": related,
                "transferabilityLevel": "High" if len(related) >= 3 else "Medium" if len(related) >= 1 else "None",
                "recommendation": f"Your {', '.join(related[:3])} experience transfers to {ms}" if related else f"No related skills found - start learning {ms} from scratch"
            })
        return details

    def get_tfidf_details(self, resume_text, job_desc):
        try:
            vectorizer = TfidfVectorizer(stop_words='english', max_features=200, ngram_range=(1, 2))
            tfidf_matrix = vectorizer.fit_transform([resume_text, job_desc])
            sim = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:2])[0][0]
            feature_names = vectorizer.get_feature_names_out()
            resume_vec = tfidf_matrix[0].toarray()[0]
            job_vec = tfidf_matrix[1].toarray()[0]
            shared_terms = []
            for i, name in enumerate(feature_names):
                if resume_vec[i] > 0 and job_vec[i] > 0:
                    combined = resume_vec[i] + job_vec[i]
                    shared_terms.append({"term": name, "relevance": round(combined, 3)})
            shared_terms.sort(key=lambda x: x['relevance'], reverse=True)
            top_resume = [{"term": name, "weight": round(resume_vec[i], 3)} for i, name in enumerate(feature_names) if resume_vec[i] > 0.1]
            top_resume.sort(key=lambda x: x['weight'], reverse=True)
            top_job = [{"term": name, "weight": round(job_vec[i], 3)} for i, name in enumerate(feature_names) if job_vec[i] > 0.1]
            top_job.sort(key=lambda x: x['weight'], reverse=True)
            return {
                "similarityPercent": round(sim * 100, 1),
                "topSharedTerms": shared_terms[:15],
                "topResumeTerms": top_resume[:10],
                "topJobTerms": top_job[:10],
                "interpretation": "Strong alignment" if sim > 0.3 else "Moderate alignment" if sim > 0.15 else "Weak alignment - tailor resume"
            }
        except:
            return {"similarityPercent": 0, "topSharedTerms": [], "topResumeTerms": [], "topJobTerms": [], "interpretation": "Could not compute"}

    def get_keyword_details(self, resume_text, job_desc):
        try:
            stop = set(stopwords.words('english'))
        except:
            stop = set()
        stop.update(['experience', 'work', 'team', 'company', 'position', 'role', 'required', 'preferred', 'ability',
                    'strong', 'knowledge', 'skills', 'years', 'working', 'including', 'using', 'etc', 'responsibilities',
                    'requirements', 'qualifications', 'job', 'must', 'will', 'also', 'well', 'good', 'great', 'new',
                    'looking', 'seeking', 'ideal', 'candidate', 'apply'])
        jw = [w for w in re.findall(r'[a-zA-Z]+', job_desc.lower()) if len(w) > 2 and w not in stop]
        unique_jw = list(dict.fromkeys(jw))
        rl = resume_text.lower()
        matched_kw = [w for w in unique_jw if w in rl]
        unmatched_kw = [w for w in unique_jw if w not in rl]
        pct = (len(matched_kw) / len(unique_jw) * 100) if unique_jw else 50
        return {
            "totalKeywords": len(unique_jw),
            "matchedCount": len(matched_kw),
            "unmatchedCount": len(unmatched_kw),
            "matchPercent": round(pct, 1),
            "matchedKeywords": matched_kw[:30],
            "unmatchedKeywords": unmatched_kw[:20],
            "suggestion": "Good keyword coverage" if pct >= 60 else "Add more JD keywords to resume" if pct >= 30 else "Resume needs significant keyword optimization"
        }

    # =======================================================================
    # EXPERIENCE EXTRACTION - ACCURATE MONTH CALCULATION (FIXED)
    # =======================================================================
    def extract_work_experiences(self, text):
        """Extract structured work experience entries with accurate month calculation"""
        experiences = []
        lines = text.split('\n')
        tl = text.lower()
        cy = datetime.now().year
        cm = datetime.now().month

        # Find WORK EXPERIENCE section boundaries
        exp_section_start = -1
        exp_section_end = len(lines)
        
        exp_section_headers = ['work experience', 'professional experience', 'employment history',
                              'work history', 'experience', 'employment', 'career history',
                              'professional background', 'relevant experience']
        non_exp_headers = ['education', 'academic', 'skills', 'technical skills', 'projects',
                          'certifications', 'achievements', 'awards', 'interests', 'hobbies',
                          'references', 'summary', 'objective', 'profile', 'about',
                          'publications', 'languages', 'volunteer', 'training', 'courses']

        for i, line in enumerate(lines):
            line_stripped = line.strip().lower()
            if len(line_stripped) > 50:
                continue
            if exp_section_start == -1:
                for header in exp_section_headers:
                    if (line_stripped == header or line_stripped.startswith(header + ':') or
                        line_stripped.startswith(header + ' -') or
                        (line_stripped.endswith(':') and header in line_stripped)):
                        if not any(neh in line_stripped for neh in ['education', 'academic', 'training']):
                            exp_section_start = i
                            break
            elif exp_section_start >= 0:
                for neh in non_exp_headers:
                    if (line_stripped == neh or line_stripped.startswith(neh + ':') or
                        line_stripped.startswith(neh + ' ') or
                        (line_stripped.endswith(':') and neh in line_stripped)):
                        if len(line_stripped) < 40:
                            exp_section_end = i
                            break
                if exp_section_end < len(lines):
                    break

        if exp_section_start >= 0:
            search_lines = lines[exp_section_start:exp_section_end]
            search_start_idx = exp_section_start
        else:
            search_lines = lines
            search_start_idx = 0

        mp = r'(?:jan(?:uary)?|feb(?:ruary)?|mar(?:ch)?|apr(?:il)?|may|jun(?:e)?|jul(?:y)?|aug(?:ust)?|sep(?:t(?:ember)?)?|oct(?:ober)?|nov(?:ember)?|dec(?:ember)?)'
        pp = r'(?:present|current|now|ongoing|till\s*date|to\s*date|today)'
        ds = r'\s*(?:[-–—]+|\bto\b)\s*'

        job_title_patterns = [
            r'(?:software|web|frontend|backend|full[\-\s]?stack|senior|junior|lead|staff|principal|associate|sr\.?|jr\.?)\s+(?:engineer|developer|programmer|architect)',
            r'(?:data|ml|machine\s+learning|ai|cloud|devops|platform|mobile|ios|android|qa|test|automation|site\s+reliability|sre)\s+(?:engineer|scientist|analyst|developer)',
            r'(?:project|product|engineering|program|delivery|technical|it)\s+manager',
            r'(?:technical|solutions?|enterprise|software|system|data)\s+architect',
            r'(?:engineering|technical|development|it)\s+(?:manager|director|lead)',
            r'(?:vp|vice\s+president|director|head)\s+(?:of\s+)?(?:engineering|technology|product|development)',
            r'(?:business|systems?|data|financial|operations?|quality)\s+analyst',
            r'(?:tech(?:nical)?|team|project|development)\s+lead',
            r'scrum\s+master',
            r'(?:devops|cloud|infrastructure|security|network)\s+(?:engineer|specialist|administrator)',
            r'(?:database|system|network|linux|windows)\s+administrator',
            r'(?:ui|ux|ui/ux|product)\s+designer',
            r'(?:senior\s+)?(?:consultant|contractor|freelancer)',
            r'(?:cto|cio|ceo|chief\s+(?:technology|technical|information)\s+officer)',
        ]

        non_job_indicators = ['student', 'course', 'coursework', 'university', 'college', 'school',
                             'degree', 'bachelor', 'master', 'phd', 'thesis', 'dissertation',
                             'gpa', 'cgpa', 'semester', 'academic', 'study', 'studied',
                             'project:', 'personal project', 'side project', 'hobby project',
                             'certificate', 'certification', 'certified', 'training',
                             'volunteer', 'volunteering', 'club', 'society', 'member of']

        for i, line in enumerate(search_lines):
            line_lower = line.lower().strip()
            original_line = line.strip()
            
            if not line_lower or len(line_lower) < 5:
                continue
            if any(ni in line_lower for ni in non_job_indicators):
                continue

            for tp in job_title_patterns:
                m = re.search(tp, line_lower)
                if m:
                    title = m.group(0).strip().title()
                    company = ""
                    dates = ""
                    duration = ""
                    start_year = None
                    end_year = None
                    start_month = None
                    end_month = None
                    total_months = 0

                    actual_idx = search_start_idx + i
                    nearby_start = max(0, actual_idx - 1)
                    nearby_end = min(len(lines), actual_idx + 4)
                    nearby_text = "\n".join(lines[nearby_start:nearby_end])
                    nearby_lower = nearby_text.lower()

                    if any(ni in nearby_lower for ni in non_job_indicators):
                        continue

                    # Extract dates with MONTH precision
                    date_m = re.search(rf'({mp})[\.,]?\s*((?:19|20)\d{{2}}){ds}(?:({mp})[\.,]?\s*((?:19|20)\d{{2}})|({pp}))', nearby_lower)
                    if date_m:
                        try:
                            start_year = int(date_m.group(2))
                            sm_name = date_m.group(1)[:3]
                            start_month = self.month_map.get(sm_name, 1)
                            
                            if date_m.group(5):
                                end_year = cy
                                end_month = cm
                                dates = f"{sm_name.title()} {start_year} - Present"
                            elif date_m.group(3) and date_m.group(4):
                                end_year = int(date_m.group(4))
                                em_name = date_m.group(3)[:3]
                                end_month = self.month_map.get(em_name, 12)
                                dates = f"{sm_name.title()} {start_year} - {em_name.title()} {end_year}"
                            
                            if start_year and end_year and start_month and end_month:
                                total_months = (end_year - start_year) * 12 + (end_month - start_month)
                                if total_months < 0:
                                    total_months = 0
                                elif total_months == 0:
                                    total_months = 1
                                
                                years = total_months // 12
                                months = total_months % 12
                                if years > 0 and months > 0:
                                    duration = f"{years}y {months}m"
                                elif years > 0:
                                    duration = f"{years}y"
                                else:
                                    duration = f"{months}m"
                        except:
                            pass
                    else:
                        year_m = re.search(r'((?:19|20)\d{2})\s*[-–—]\s*(?:((?:19|20)\d{2})|(' + pp + r'))', nearby_lower)
                        if year_m:
                            try:
                                start_year = int(year_m.group(1))
                                start_month = 1
                                if year_m.group(3):
                                    end_year = cy
                                    end_month = cm
                                    dates = f"{start_year} - Present"
                                elif year_m.group(2):
                                    end_year = int(year_m.group(2))
                                    end_month = 12
                                    dates = f"{start_year} - {end_year}"
                                
                                if start_year and end_year:
                                    total_months = (end_year - start_year) * 12 + (end_month - start_month)
                                    if total_months <= 0:
                                        total_months = 6
                                    years = total_months // 12
                                    months = total_months % 12
                                    if years > 0 and months > 0:
                                        duration = f"{years}y {months}m"
                                    elif years > 0:
                                        duration = f"{years}y"
                                    else:
                                        duration = f"{months}m"
                            except:
                                pass

                    # Extract company
                    company_patterns = [
                        r'(?:at|@)\s+([A-Z][A-Za-z0-9\s&.,]+?)(?:\s*[,\|\n]|\s+(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec|from|\d{4}))',
                        r'([A-Z][A-Za-z0-9\s&]+(?:Inc\.?|LLC|Ltd\.?|Corp\.?|Technologies|Solutions|Software|Systems|Services|Labs|Studio))',
                        r'[-–—,]\s*([A-Z][A-Za-z0-9\s&.,]+?)(?:\s*[,\|\n]|\s*$)',
                    ]
                    for cp in company_patterns:
                        cm_match = re.search(cp, original_line + " " + (lines[min(actual_idx + 1, len(lines) - 1)] if actual_idx + 1 < len(lines) else ""))
                        if cm_match:
                            company_candidate = cm_match.group(1).strip()
                            company_candidate = re.sub(r'\s+', ' ', company_candidate).strip('.,;:- ')
                            if (len(company_candidate) >= 2 and len(company_candidate) <= 60 and
                                not any(ni in company_candidate.lower() for ni in non_job_indicators)):
                                company = company_candidate
                                break

                    # Extract responsibilities
                    desc_lines = []
                    for j in range(i + 1, min(i + 8, len(search_lines))):
                        l = search_lines[j].strip()
                        if not l:
                            continue
                        if l.startswith(('•', '-', '*', '–', '▪', '►', '○')) or (len(l) > 20 and l[0].isupper()):
                            l_lower = l.lower()
                            if any(re.search(tp2, l_lower) for tp2 in job_title_patterns):
                                break
                            if any(neh in l_lower for neh in non_exp_headers):
                                break
                            desc_lines.append(l.lstrip('•-*–▪►○ ')[:200])
                        elif len(l) < 30 and any(re.search(tp2, l.lower()) for tp2 in job_title_patterns):
                            break
                        if len(desc_lines) >= 5:
                            break

                    if dates or company:
                        already = any(e['title'].lower() == title.lower() and
                                     (e.get('company', '').lower() == company.lower() or e.get('dates', '') == dates)
                                     for e in experiences)
                        if not already:
                            experiences.append({
                                "title": title,
                                "company": company or "Not specified",
                                "dates": dates or "Not specified",
                                "duration": duration,
                                "startYear": start_year,
                                "endYear": end_year,
                                "startMonth": start_month,
                                "endMonth": end_month,
                                "totalMonths": total_months,
                                "responsibilities": desc_lines[:5],
                                "foundInSection": "Work Experience section" if exp_section_start >= 0 else "Full resume scan",
                            })
                    break

        experiences.sort(key=lambda x: (x.get('startYear') or 0, x.get('startMonth') or 0), reverse=True)
        return experiences

    def extract_experience(self, text):
        """Extract total years of experience with ACCURATE month calculation"""
        tl = text.lower()
        
        # Method 1: Explicit statement
        explicit_patterns = [
            r'(\d+)\+?\s*(?:years?|yrs?)\s*(?:and\s*)?(\d+)?\s*(?:months?|mos?)?\s*(?:of\s+)?(?:total\s+)?(?:professional\s+)?experience',
            r'(?:total|overall)\s*(?:of\s+)?(\d+)\+?\s*(?:years?|yrs?)\s*(?:and\s*)?(\d+)?\s*(?:months?|mos?)?\s*(?:of\s+)?experience',
            r'(\d+)\+?\s*(?:years?|yrs?)\s*(?:and\s*)?(\d+)?\s*(?:months?|mos?)?\s*(?:of\s+)?(?:professional|work|industry|hands[\-\s]?on|relevant)\s*experience',
        ]
        
        for p in explicit_patterns:
            m = re.search(p, tl)
            if m:
                try:
                    years = int(m.group(1))
                    months = int(m.group(2)) if m.lastindex >= 2 and m.group(2) else 0
                    total_months = years * 12 + months
                    if 0 < total_months < 600:
                        ctx_start = max(0, m.start() - 100)
                        ctx_end = min(len(tl), m.end() + 50)
                        ctx = tl[ctx_start:ctx_end]
                        if not any(edu in ctx for edu in ['education', 'university', 'college', 'degree', 'academic', 'gpa']):
                            return total_months / 12
                except:
                    pass

        # Method 2: Calculate from work experiences (ACCURATE)
        experiences = self.extract_work_experiences(text)
        
        if experiences:
            total_months = 0
            for exp in experiences:
                exp_months = exp.get('totalMonths', 0)
                if exp_months > 0:
                    total_months += exp_months
                elif exp.get('startYear') and exp.get('endYear'):
                    sy = exp.get('startYear')
                    ey = exp.get('endYear')
                    sm = exp.get('startMonth', 1)
                    em = exp.get('endMonth', 12)
                    months = (ey - sy) * 12 + (em - sm)
                    if months > 0:
                        total_months += months
            
            if total_months > 0:
                return total_months / 12

        # Method 3: Date range calculation
        yfd = self._calc_dates_strict(text)
        if yfd > 0:
            return yfd

        return 0

    def _calc_dates_strict(self, text):
        """Calculate experience from dates with ACCURATE month calculation"""
        lines = text.split('\n')
        tl = text.lower()
        cy = datetime.now().year
        cm = datetime.now().month

        exp_section_start = -1
        exp_section_end = len(lines)
        
        exp_headers = ['work experience', 'professional experience', 'employment', 'experience', 'career']
        non_exp_headers = ['education', 'academic', 'skills', 'projects', 'certifications', 'training']

        for i, line in enumerate(lines):
            ll = line.strip().lower()
            if len(ll) > 50:
                continue
            if exp_section_start == -1:
                for h in exp_headers:
                    if ll == h or ll.startswith(h + ':') or ll.startswith(h + ' '):
                        if 'education' not in ll and 'academic' not in ll:
                            exp_section_start = i
                            break
            elif exp_section_start >= 0:
                for nh in non_exp_headers:
                    if ll == nh or ll.startswith(nh + ':'):
                        exp_section_end = i
                        break
                if exp_section_end < len(lines):
                    break

        if exp_section_start == -1:
            return 0

        exp_text = '\n'.join(lines[exp_section_start:exp_section_end]).lower()

        mp = r'(?:jan(?:uary)?|feb(?:ruary)?|mar(?:ch)?|apr(?:il)?|may|jun(?:e)?|jul(?:y)?|aug(?:ust)?|sep(?:t(?:ember)?)?|oct(?:ober)?|nov(?:ember)?|dec(?:ember)?)'
        pp = r'(?:present|current|now|ongoing|till\s*date|to\s*date|today)'
        ds = r'\s*(?:[-–—]+|\bto\b)\s*'

        ranges = []
        
        for m in re.finditer(rf'({mp})[\.,]?\s*((?:19|20)\d{{2}}){ds}(?:({mp})[\.,]?\s*((?:19|20)\d{{2}})|({pp}))', exp_text):
            try:
                sy = int(m.group(2))
                sm = self.month_map.get(m.group(1)[:3], 1)
                if m.group(5):
                    ey, em = cy, cm
                else:
                    ey = int(m.group(4))
                    em = self.month_map.get(m.group(3)[:3], 12) if m.group(3) else 12
                if 1980 <= sy <= cy and sy <= ey <= cy + 1:
                    ranges.append((sy, sm, ey, em))
            except:
                pass

        if ranges:
            tm = self._merge_ranges(ranges)
            if tm >= 1:
                return tm / 12
        return 0

    def _merge_ranges(self, ranges):
        """Merge overlapping date ranges and calculate total months ACCURATELY"""
        if not ranges:
            return 0
        month_ranges = []
        for sy, sm, ey, em in ranges:
            start = sy * 12 + sm
            end = ey * 12 + em
            if end > start:
                month_ranges.append((start, end))
        if not month_ranges:
            return 0
        month_ranges.sort()
        merged = [month_ranges[0]]
        for start, end in month_ranges[1:]:
            if start <= merged[-1][1]:
                merged[-1] = (merged[-1][0], max(merged[-1][1], end))
            else:
                merged.append((start, end))
        return sum(end - start for start, end in merged)

    def get_experience_details(self, resume_text, job_desc):
        """Get detailed experience breakdown with ACCURATE calculations"""
        experiences = self.extract_work_experiences(resume_text)
        total_years_decimal = self.extract_experience(resume_text)
        
        total_years = int(total_years_decimal)
        total_months = round((total_years_decimal - total_years) * 12)
        
        if total_years > 0 and total_months > 0:
            experience_display = f"{total_years}y {total_months}m"
        elif total_years > 0:
            experience_display = f"{total_years}y"
        elif total_months > 0:
            experience_display = f"{total_months}m"
        else:
            experience_display = "0"
        
        req_exp = 0
        req_patterns = [
            r'(\d+)\+?\s*(?:years?|yrs?)\s*(?:of\s+)?(?:experience|exp)',
            r'(?:minimum|at\s+least|require[sd]?)\s*(\d+)\+?\s*(?:years?|yrs?)',
        ]
        for p in req_patterns:
            m = re.search(p, job_desc.lower())
            if m:
                try:
                    req_exp = int(m.group(1))
                    break
                except:
                    pass

        if total_years_decimal > 0:
            detection_method = "Calculated from work history" if experiences else "Explicit statement in resume"
        else:
            detection_method = "No experience detected"

        calculated_months = sum(exp.get('totalMonths', 0) for exp in experiences)
        calculated_years_decimal = calculated_months / 12 if calculated_months > 0 else 0

        if req_exp > 0:
            meets_req = total_years_decimal >= req_exp
            gap = max(0, req_exp - total_years_decimal)
            gap_years = int(gap)
            gap_months = int((gap - gap_years) * 12)
            gap_display = f"{gap_years}y {gap_months}m" if gap_years > 0 and gap_months > 0 else f"{gap_years}y" if gap_years > 0 else f"{gap_months}m" if gap_months > 0 else "0"
        else:
            meets_req = True
            gap = 0
            gap_display = "0"

        tips = []
        if total_years_decimal == 0:
            tips.append("⚠️ No work experience detected — add a clear 'Work Experience' section")
            tips.append("💡 Format: Job Title | Company | Dates (e.g., 'Jan 2023 - Apr 2024')")
        elif not experiences:
            tips.append("⚠️ Experience mentioned but couldn't extract details")
        if req_exp > 0 and not meets_req:
            tips.append(f"⚠️ JD requires {req_exp}y, you have {experience_display} ({gap_display} gap)")
        if not tips:
            tips.append("✅ Experience section parsed successfully")

        return {
            "description": "Work experience detected from your resume",
            "totalYears": round(total_years_decimal, 2),
            "totalYearsDisplay": experience_display,
            "totalMonths": round(total_years_decimal * 12),
            "calculatedYears": round(calculated_years_decimal, 2),
            "requiredYears": req_exp,
            "meetsRequirement": meets_req,
            "experienceGap": round(gap, 2),
            "experienceGapDisplay": gap_display,
            "positionsFound": len(experiences),
            "detectionMethod": detection_method,
            "positions": experiences,
            "tips": tips,
            "recommendedFormat": "Job Title | Company Name\nMonth Year - Month Year (or Present)\n• Achievement with metrics",
        }

    # =======================================================================
    # EDUCATION EXTRACTION - FIXED VERSION
    # =======================================================================
    def extract_education_details(self, text):
        """Extract structured education entries with strict validation"""
        education = []
        tl = text.lower()
        lines = text.split('\n')

        edu_section_start = -1
        edu_section_end = len(text)
        edu_section_headers = ['education', 'academic background', 'academic qualifications',
                              'educational qualifications', 'academic details', 'academics',
                              'educational background', 'qualifications', 'schooling']
        non_edu_headers = ['experience', 'employment', 'work history', 'professional experience',
                          'skills', 'technical skills', 'projects', 'certifications',
                          'achievements', 'awards', 'interests', 'hobbies', 'references',
                          'summary', 'objective', 'profile', 'about', 'contact',
                          'publications', 'languages', 'volunteer']

        for i, line in enumerate(lines):
            line_stripped = line.strip().lower()
            if edu_section_start == -1:
                for header in edu_section_headers:
                    if (line_stripped == header or line_stripped.startswith(header + ':') or
                        line_stripped.startswith(header + ' ') or
                        (len(line_stripped) < 40 and header in line_stripped and
                         not any(neh in line_stripped for neh in ['experience', 'skills']))):
                        edu_section_start = sum(len(l) + 1 for l in lines[:i])
                        break
            elif edu_section_start >= 0:
                for neh in non_edu_headers:
                    if (line_stripped == neh or line_stripped.startswith(neh + ':') or line_stripped.startswith(neh + ' ')):
                        if len(line_stripped) < 50:
                            edu_section_end = sum(len(l) + 1 for l in lines[:i])
                            break
                if edu_section_end < len(text):
                    break

        if edu_section_start >= 0:
            edu_text = text[edu_section_start:edu_section_end]
            edu_text_lower = edu_text.lower()
        else:
            edu_text = text
            edu_text_lower = tl

        degree_patterns = [
            (r'\b(?:ph\.?d\.?|doctorate|doctoral)\b(?:\s+(?:in|of)\s+)?([A-Za-z\s&]{2,50})?', 'PhD'),
            (r"\b(?:master'?s?\s*(?:degree|of)?|m\.?s\.?\b(?!\w)|m\.?sc\.?\b|m\.?b\.?a\.?\b|m\.?tech\.?\b|m\.?e\.?\b(?!\w)|m\.?c\.?a\.?\b)\s*(?:in|of)?\s*([A-Za-z\s&]{2,50})?", 'Masters'),
            (r"\b(?:bachelor'?s?\s*(?:degree|of)?|b\.?s\.?\b(?!\w)|b\.?sc\.?\b|b\.?tech\.?\b|b\.?e\.?\b(?!\w)|b\.?a\.?\b(?!\w)|b\.?c\.?a\.?\b)\s*(?:in|of)?\s*([A-Za-z\s&]{2,50})?", 'Bachelors'),
            (r'\b(?:diploma|associate\s+degree|associate\s+of)\b\s*(?:in|of)?\s*([A-Za-z\s&]{2,50})?', 'Diploma'),
        ]

        field_stopwords = {'the', 'and', 'from', 'with', 'at', 'by', 'for', 'on', 'to', 'is', 'was', 'are', 'were',
                          'has', 'have', 'had', 'been', 'will', 'would', 'could', 'should', 'may', 'might', 'can',
                          'a', 'an', 'this', 'that', 'these', 'those', 'it', 'its', 'i', 'we', 'they', 'he', 'she',
                          'my', 'our', 'your', 'their', 'experience', 'work', 'working', 'worked', 'project', 'team',
                          'developed', 'managed', 'built', 'created', 'designed', 'used', 'using', 'responsible',
                          'skills', 'proficient', 'knowledge', 'university', 'college', 'institute', 'school', 'academy'}

        def clean_field(raw_field):
            if not raw_field:
                return ""
            words = raw_field.strip().split()
            cleaned = []
            for w in words:
                wl = w.lower().strip('.,;:-')
                if wl in field_stopwords:
                    if cleaned:
                        break
                    continue
                if len(wl) < 2:
                    continue
                cleaned.append(w)
                if len(cleaned) >= 6:
                    break
            result = ' '.join(cleaned).strip().title()
            return "" if len(result) < 3 or result.lower() in field_stopwords else result

        for pattern, deg_type in degree_patterns:
            for m in re.finditer(pattern, edu_text_lower, re.IGNORECASE):
                raw_field = m.group(1) if m.lastindex and m.group(1) else ""
                field = clean_field(raw_field)
                ctx_start = max(0, m.start() - 200)
                ctx_end = min(len(edu_text), m.end() + 200)
                ctx = edu_text[ctx_start:ctx_end]
                ctx_lower = ctx.lower()

                institution = ""
                inst_patterns = [
                    r'([A-Z][A-Za-z\s\.]+(?:University|College|Institute|School|Academy))',
                    r'(?:University|College|Institute|School|Academy)\s+(?:of\s+)?([A-Z][A-Za-z\s]+)',
                ]
                for ip in inst_patterns:
                    im = re.search(ip, ctx)
                    if im:
                        inst_candidate = im.group(1).strip() if im.group(1) else im.group(0).strip()
                        inst_candidate = inst_candidate.strip('.,;:- ')
                        if len(inst_candidate) >= 5:
                            institution = inst_candidate[:80].title()
                            break

                year = ""
                for yp in [r'(?:graduated|graduation|passing|completed|class\s+of)\s*(?:in\s+)?(\d{4})',
                          r'(\d{4})\s*[-–]\s*(\d{4})', r'\b(20[0-2]\d|19[89]\d)\b']:
                    ym = re.search(yp, ctx_lower)
                    if ym:
                        year = ym.group(2) if ym.lastindex == 2 else ym.group(1)
                        try:
                            if int(year) < 1970 or int(year) > datetime.now().year + 6:
                                year = ""
                                continue
                        except:
                            year = ""
                            continue
                        break

                gpa = ""
                for gp in [r'(?:gpa|cgpa|cpi|grade)[:\s]*(\d+\.?\d*)\s*(?:/\s*(\d+\.?\d*))?']:
                    gm = re.search(gp, ctx_lower)
                    if gm:
                        gpa_val = gm.group(1)
                        gpa_max = gm.group(2) if gm.lastindex >= 2 and gm.group(2) else ""
                        try:
                            gv = float(gpa_val)
                            if gpa_max:
                                gpa = f"{gpa_val}/{gpa_max}"
                            elif gv <= 4.0:
                                gpa = f"{gpa_val}/4.0"
                            elif gv <= 10.0:
                                gpa = f"{gpa_val}/10.0"
                            elif gv <= 100:
                                gpa = f"{gpa_val}%"
                        except:
                            pass
                        break

                snippet = edu_text[max(0, m.start() - 20):min(len(edu_text), m.end() + 100)].strip().replace('\n', ' ')
                snippet = re.sub(r'\s+', ' ', snippet)[:150]

                already = any(e['degreeType'] == deg_type and (e.get('institution', '') == institution or e.get('field', '') == field) for e in education)
                if not already:
                    education.append({
                        "degreeType": deg_type,
                        "field": field if field and len(field) > 2 else "Not specified",
                        "institution": institution or "Not specified",
                        "year": year,
                        "gpa": gpa,
                        "resumeSnippet": f"...{snippet}...",
                        "foundInSection": "Education section" if edu_section_start >= 0 else "Full resume scan",
                    })

        if not education:
            quick_checks = [
                (r'\bB\.?Tech\b', 'Bachelors', 'Technology'), (r'\bM\.?Tech\b', 'Masters', 'Technology'),
                (r'\bBCA\b', 'Bachelors', 'Computer Applications'), (r'\bMCA\b', 'Masters', 'Computer Applications'),
                (r'\bBBA\b', 'Bachelors', 'Business Administration'), (r'\bMBA\b', 'Masters', 'Business Administration'),
                (r'\bB\.?Sc\b', 'Bachelors', 'Science'), (r'\bM\.?Sc\b', 'Masters', 'Science'),
            ]
            for pat, deg, default_field in quick_checks:
                qm = re.search(pat, edu_text, re.IGNORECASE)
                if qm:
                    ctx = edu_text[max(0, qm.start() - 150):min(len(edu_text), qm.end() + 150)]
                    snippet = edu_text[max(0, qm.start() - 20):min(len(edu_text), qm.end() + 100)].strip().replace('\n', ' ')
                    institution = ""
                    im = re.search(r'([A-Z][A-Za-z\s\.]+(?:University|College|Institute))', ctx)
                    if im:
                        institution = im.group(1).strip()[:80].title()
                    year = ""
                    ym = re.search(r'\b(20[0-2]\d|19[89]\d)\b', ctx)
                    if ym:
                        year = ym.group(1)
                    if not any(e['degreeType'] == deg for e in education):
                        education.append({
                            "degreeType": deg, "field": default_field,
                            "institution": institution or "Not specified", "year": year, "gpa": "",
                            "resumeSnippet": f"...{re.sub(r' +', ' ', snippet)[:150]}...",
                            "foundInSection": "Education section" if edu_section_start >= 0 else "Full resume scan",
                        })

        return education

    def extract_education(self, text):
        tl = text.lower()
        if any(w in tl for w in ['phd', 'ph.d', 'doctorate', 'doctoral']):
            return 15
        elif any(w in tl for w in ["master's", 'masters', 'msc', 'm.sc', 'mba', 'ms degree', 'master of', 'm.tech', 'mtech']):
            return 12
        elif any(w in tl for w in ["bachelor's", 'bachelors', 'bsc', 'b.sc', 'b.tech', 'btech', 'b.e.', 'bachelor of']):
            return 10
        elif any(w in tl for w in ['diploma', 'associate', 'certification', 'bootcamp']):
            return 7
        elif any(w in tl for w in ['degree', 'university', 'college', 'institute']):
            return 8
        return 0

    # =======================================================================
    # OTHER EXTRACTION METHODS
    # =======================================================================
    def detect_certifications_detailed(self, text):
        tl = text.lower()
        found = []
        for cert in KNOWN_CERTIFICATIONS:
            if cert in tl:
                if not any(cert in fc['name'] or fc['name'] in cert for fc in found):
                    idx = tl.find(cert)
                    ctx_s = max(0, idx - 40)
                    ctx_e = min(len(tl), idx + len(cert) + 40)
                    snippet = text[ctx_s:ctx_e].strip().replace('\n', ' ')
                    found.append({"name": cert.title(), "context": f"...{snippet}..."})
        return found

    def detect_achievements_detailed(self, text):
        achievements = []
        patterns = [
            (r'(?:increased|improved|boosted|grew|raised|enhanced)\s+.*?\d+\s*%', 'percentage_increase'),
            (r'(?:reduced|decreased|lowered|cut|minimized|saved)\s+.*?\d+\s*%', 'percentage_decrease'),
            (r'\$\s*\d+[\d,]*(?:\.\d+)?\s*(?:million|billion|m|b|k)?', 'revenue_money'),
            (r'\d+[\d,]*\s*(?:users|customers|clients|visitors|downloads)', 'user_metrics'),
            (r'(?:managed|led|mentored|supervised)\s+(?:a\s+)?(?:team\s+of\s+)?\d+', 'team_size'),
            (r'\d+[%]\s*(?:increase|decrease|improvement|reduction|growth)', 'metric_change'),
            (r'(?:awarded|award|recognition|prize|winner|patent)', 'award'),
        ]
        tl = text.lower()
        for p, category in patterns:
            for m in re.finditer(p, tl):
                snippet = m.group(0).strip()
                if snippet and snippet not in [a['text'] for a in achievements]:
                    idx = m.start()
                    line_start = text.rfind('\n', 0, idx) + 1
                    line_end = text.find('\n', idx)
                    if line_end == -1:
                        line_end = min(idx + 150, len(text))
                    full_line = text[line_start:line_end].strip()
                    achievements.append({"text": snippet, "category": category.replace('_', ' ').title(), "fullContext": full_line[:200]})
        return achievements

    def get_quality_details(self, text, skills):
        word_count = len(text.split())
        char_count = len(text)
        line_count = len(text.split('\n'))
        sections_detected = []
        section_keywords = {
            'Summary/Objective': ['summary', 'objective', 'profile', 'about me'],
            'Experience': ['experience', 'employment', 'work history'],
            'Education': ['education', 'academic', 'qualification', 'degree'],
            'Skills': ['skills', 'technical skills', 'technologies'],
            'Projects': ['projects', 'portfolio'],
            'Certifications': ['certifications', 'certificates'],
            'Contact': ['email', 'phone', 'linkedin', 'github'],
        }
        tl = text.lower()
        for section, keywords in section_keywords.items():
            if any(kw in tl for kw in keywords):
                sections_detected.append(section)
        has_bullet_points = bool(re.search(r'[•\-\*]', text))
        has_numbers = bool(re.search(r'\d+[%$]|\$\d+|\d+\s*(?:users|projects|years)', tl))
        return {
            "wordCount": word_count, "characterCount": char_count, "lineCount": line_count,
            "skillsDetected": len(skills), "sectionsFound": sections_detected,
            "sectionCount": len(sections_detected), "hasBulletPoints": has_bullet_points,
            "hasQuantifiedResults": has_numbers,
            "lengthAssessment": "Good length" if 400 <= word_count <= 1200 else "Too short" if word_count < 400 else "Consider trimming",
            "tips": [
                "Add more skills" if len(skills) < 8 else "Good skill coverage",
                "Add bullet points" if not has_bullet_points else "Good use of bullet points",
                "Add numbers/metrics" if not has_numbers else "Good use of metrics",
            ]
        }

    def get_seniority_details(self, resume_text, job_desc):
        def detect_level(text):
            tl = text.lower()
            if any(w in tl for w in ['principal engineer', 'staff engineer', 'cto', 'chief technology', 'head of engineering', 'vp of engineering']):
                return 'principal'
            elif any(w in tl for w in ['senior engineer', 'senior developer', 'lead engineer', 'tech lead', 'team lead', 'sr.', 'sr ']):
                return 'senior'
            elif any(w in tl for w in ['mid-level', 'mid level', 'software engineer ii', 'developer ii']):
                return 'mid'
            elif any(w in tl for w in ['junior', 'associate', 'entry-level', 'entry level', 'software engineer i', 'graduate']):
                return 'junior'
            elif any(w in tl for w in ['intern', 'trainee', 'apprentice']):
                return 'intern'
            return 'unknown'

        c_level = detect_level(resume_text)
        j_level = detect_level(job_desc)
        level_names = {'intern': 'Intern', 'junior': 'Junior', 'mid': 'Mid-Level', 'senior': 'Senior', 'principal': 'Principal/Staff', 'unknown': 'Not Specified'}
        level_order = {'intern': 0, 'junior': 1, 'mid': 2, 'senior': 3, 'principal': 4, 'unknown': -1}
        c_ord = level_order.get(c_level, -1)
        j_ord = level_order.get(j_level, -1)

        if j_ord == -1 or c_ord == -1:
            match_desc, match_quality = "Could not determine levels", "Neutral"
        elif c_ord == j_ord:
            match_desc, match_quality = "Perfect seniority match", "Excellent"
        elif c_ord == j_ord + 1:
            match_desc, match_quality = "Overqualified but strong", "Good"
        elif c_ord == j_ord - 1:
            match_desc, match_quality = "Stretch role", "Moderate"
        elif c_ord > j_ord:
            match_desc, match_quality = "Significantly overqualified", "Concern"
        else:
            match_desc, match_quality = "Underqualified", "Gap"

        return {
            "candidateLevel": level_names.get(c_level, 'Unknown'),
            "jobLevel": level_names.get(j_level, 'Unknown'),
            "candidateLevelRaw": c_level, "jobLevelRaw": j_level,
            "matchDescription": match_desc, "matchQuality": match_quality,
        }

    def find_related_skills(self, missing, resume_skills):
        ml = missing.lower()
        rl = [s.lower() for s in resume_skills]
        c = 0
        for gs in RELATED_SKILL_GROUPS.values():
            if ml in gs:
                c += sum(1 for rs in rl if rs in gs and rs != ml)
        return c

    # =======================================================================
    # MAIN SCORING FUNCTION
    # =======================================================================
    def calculate_score(self, matched, missing, resume_text, job_desc, exp_years_input, resume_skills):
        total = len(matched) + len(missing)

        # 1. Core Skill Match (25)
        skill_pct = len(matched) / total if total > 0 else 0
        skill_score = skill_pct * 25 if total > 0 else 12
        skill_details = self.get_skill_match_details(matched + missing, resume_text)

        # 2. Skill Depth (8)
        depth_details_list, avg_mentions = self.get_skill_depth_details(matched, resume_text)
        if not depth_details_list:
            depth_score = 0
        else:
            avg_depth = sum(float(d['depthScore'].split('/')[0]) for d in depth_details_list) / len(depth_details_list)
            depth_score = min(round((avg_depth / 3) * 8, 1), 8)

        # 3. Transferable (7)
        transferable_details = self.get_transferable_details(missing, resume_skills)
        related_bonus = 0
        if missing:
            tr = sum(min(self.find_related_skills(m, resume_skills), 3) for m in missing)
            mp = len(missing) * 2
            if mp > 0:
                related_bonus = min((tr / mp) * 7, 7)

        # 4. TF-IDF Cosine (10)
        tfidf_details = self.get_tfidf_details(resume_text, job_desc)
        cosine_score = min((tfidf_details['similarityPercent'] / 100) * 10, 10)

        # 5. Keywords (5)
        keyword_details = self.get_keyword_details(resume_text, job_desc)
        keyword_score = (keyword_details['matchPercent'] / 100) * 5

        # 6. Experience (15) - ACCURATE SCORING WITH DECIMAL YEARS
        experience_details = self.get_experience_details(resume_text, job_desc)
        exp_years = experience_details['totalYears']
        req_exp = experience_details['requiredYears']
        experience_entries = experience_details['positions']

        if exp_years == 0:
            exp_score = 0
        elif req_exp > 0:
            ratio = exp_years / req_exp
            if ratio >= 1.0:
                exp_score = 15
            elif ratio >= 0.75:
                exp_score = 12
            elif ratio >= 0.5:
                exp_score = 9
            elif ratio >= 0.25:
                exp_score = 6
            else:
                exp_score = 3
        else:
            if exp_years >= 10:
                exp_score = 15
            elif exp_years >= 7:
                exp_score = 13
            elif exp_years >= 5:
                exp_score = 11
            elif exp_years >= 3:
                exp_score = 9
            elif exp_years >= 2:
                exp_score = 7
            elif exp_years >= 1:
                exp_score = 5
            elif exp_years >= 0.5:
                exp_score = 3
            else:
                exp_score = 0

        # 7. Education (7)
        edu_raw = self.extract_education(resume_text)
        edu_score = min(round(edu_raw * 0.7), 7)
        education_entries = self.extract_education_details(resume_text)

        # 8. Certifications (5)
        cert_entries = self.detect_certifications_detailed(resume_text)
        cert_score = 5 if len(cert_entries) >= 4 else 4 if len(cert_entries) >= 3 else 3 if len(cert_entries) >= 2 else 2 if len(cert_entries) >= 1 else 0

        # 9. Achievements (8)
        ach_entries = self.detect_achievements_detailed(resume_text)
        ach_score = 8 if len(ach_entries) >= 6 else 6 if len(ach_entries) >= 4 else 5 if len(ach_entries) >= 3 else 4 if len(ach_entries) >= 2 else 2 if len(ach_entries) >= 1 else 0

        # 10. Quality (5)
        quality_details = self.get_quality_details(resume_text, resume_skills)
        rl = len(resume_text)
        quality = 5 if rl >= 3000 else 4 if rl >= 2000 else 3 if rl >= 1000 else 2 if rl >= 500 else 1
        if len(resume_skills) >= 15:
            quality = min(quality + 1, 5)

        # 11. Seniority (5)
        seniority_details = self.get_seniority_details(resume_text, job_desc)
        level_order = {'intern': 0, 'junior': 1, 'mid': 2, 'senior': 3, 'principal': 4, 'unknown': -1}
        c_lvl = level_order.get(seniority_details['candidateLevelRaw'], -1)
        j_lvl = level_order.get(seniority_details['jobLevelRaw'], -1)
        if j_lvl == -1 or c_lvl == -1:
            seniority_score = 3
        elif c_lvl == j_lvl:
            seniority_score = 5
        elif c_lvl == j_lvl + 1:
            seniority_score = 4
        elif c_lvl == j_lvl - 1:
            seniority_score = 3
        elif c_lvl > j_lvl:
            seniority_score = 3
        else:
            seniority_score = max(0, 5 - (j_lvl - c_lvl) * 2)

        overall = max(5, min(98, round(
            skill_score + depth_score + related_bonus + cosine_score + keyword_score +
            exp_score + edu_score + cert_score + ach_score + quality + seniority_score
        )))

        return overall, {
            "total": overall, "totalMax": 100,
            "coreSkillMatch": round(skill_score, 1), "coreSkillMatchMax": 25,
            "skillDepth": round(depth_score, 1), "skillDepthMax": 8,
            "transferableSkills": round(related_bonus, 1), "transferableSkillsMax": 7,
            "cosineSimilarity": round(cosine_score, 1), "cosineSimilarityMax": 10,
            "keywordOverlap": round(keyword_score, 1), "keywordOverlapMax": 5,
            "experience": exp_score, "experienceMax": 15,
            "education": edu_score, "educationMax": 7,
            "certifications": cert_score, "certificationsMax": 5,
            "achievements": ach_score, "achievementsMax": 8,
            "resumeQuality": round(quality, 1), "resumeQualityMax": 5,
            "seniorityMatch": seniority_score, "seniorityMatchMax": 5,
            "skillMatchPct": round(skill_pct * 100, 1),
            "keywordOverlapPct": round(keyword_details['matchPercent'], 1),
            "cosineSimilarityPct": round(tfidf_details['similarityPercent'], 1),
            "candidateExperience": exp_years, "requiredExperience": req_exp,
            "details": {
                "coreSkillMatch": {"description": "Which required skills were found", "matchedCount": len(matched), "missingCount": len(missing), "totalRequired": total, "skills": skill_details},
                "skillDepth": {"description": "How deeply each skill is demonstrated", "averageMentions": avg_mentions, "skills": depth_details_list},
                "transferableSkills": {"description": "Related skills for missing ones", "skills": transferable_details},
                "cosineSimilarity": {"description": "Semantic similarity", **tfidf_details},
                "keywordOverlap": {"description": "Direct keyword matches", **keyword_details},
                "experience": {**experience_details, "scoreAwarded": exp_score, "maxScore": 15,
                              "scoringBreakdown": f"No experience (0/15)" if exp_years == 0 else f"{experience_details['totalYearsDisplay']} vs {req_exp}y required → {exp_score}/15" if req_exp > 0 else f"{experience_details['totalYearsDisplay']} → {exp_score}/15"},
                "education": {"description": "Educational qualifications", "rawScore": edu_raw, "scoreAwarded": edu_score, "maxScore": 7, "entries": education_entries, "totalFound": len(education_entries),
                             "highestDegree": education_entries[0]['degreeType'] if education_entries else "Not detected",
                             "educationSectionFound": any(e.get('foundInSection') == 'Education section' for e in education_entries),
                             "scoringBreakdown": "PhD (7/7)" if edu_raw >= 15 else "Masters (7/7)" if edu_raw >= 12 else "Bachelors (7/7)" if edu_raw >= 10 else "Degree detected" if edu_raw >= 7 else "None detected",
                             "tips": ["✅ Education found" if education_entries else "⚠️ Add education section"]},
                "certifications": {"description": "Certifications", "count": len(cert_entries), "entries": cert_entries,
                                  "suggestion": "Great!" if len(cert_entries) >= 3 else "Consider certifications"},
                "achievements": {"description": "Quantified achievements", "count": len(ach_entries), "entries": ach_entries,
                                "suggestion": "Excellent metrics!" if len(ach_entries) >= 4 else "Add more numbers"},
                "resumeQuality": {"description": "Resume structure", **quality_details},
                "seniorityMatch": {"description": "Seniority comparison", **seniority_details},
            }
        }

    def generate_analysis(self, matched, missing, score, exp, company, job_data, sb):
        total = len(matched) + len(missing)
        pct = (len(matched) / total * 100) if total > 0 else 0
        level = "Senior" if exp >= 7 else "Mid-Level" if exp >= 4 else "Junior" if exp >= 1 else "Entry-Level"

        if score >= 80:
            rec, rr = "STRONG HIRE", f"Excellent ({score}/100)"
        elif score >= 65:
            rec, rr = "HIRE", f"Good ({score}/100)"
        elif score >= 50:
            rec, rr = "MAYBE", f"Moderate ({score}/100)"
        elif score >= 35:
            rec, rr = "CONDITIONAL", f"Below average ({score}/100)"
        else:
            rec, rr = "NEEDS DEVELOPMENT", f"Low ({score}/100)"

        summary = f"{score}/100 - {'Excellent!' if score >= 80 else 'Good' if score >= 65 else 'Decent' if score >= 50 else 'Partial match' if score >= 35 else 'Needs work'}"
        fit = f"{len(matched)}/{total} skills ({int(pct)}% match). "
        if matched:
            fit += f"Strong: {', '.join(matched[:5])}. "
        if missing:
            fit += f"Needs: {', '.join(missing[:4])}. "

        # Format experience for display
        exp_years = int(exp)
        exp_months = round((exp - exp_years) * 12)
        if exp_years > 0 and exp_months > 0:
            exp_display = f"{exp_years}y {exp_months}m"
        elif exp_years > 0:
            exp_display = f"{exp_years}y"
        elif exp_months > 0:
            exp_display = f"{exp_months}m"
        else:
            exp_display = "0"
        
        fit += f"{exp_display} exp ({level}). Score: {score}/100."

        strengths = []
        if matched:
            strengths.append(f"Proficient in {len(matched)} required skills: {', '.join(matched[:6])}")
        if exp >= 5:
            strengths.append(f"Strong: {exp_display} experience")
        elif exp >= 1:
            strengths.append(f"{exp_display} experience")
        certs = sb.get('details', {}).get('certifications', {}).get('count', 0)
        if certs >= 2:
            strengths.append(f"{certs} relevant certifications")
        achs = sb.get('details', {}).get('achievements', {}).get('count', 0)
        if achs >= 3:
            strengths.append(f"{achs} quantified achievements")
        if sb.get('cosineSimilarityPct', 0) >= 30:
            strengths.append("Resume aligns well with JD")
        if not strengths:
            strengths.append("Shows willingness to learn")

        weaknesses = []
        if missing:
            weaknesses.append(f"Missing {len(missing)} skills: {', '.join(missing[:4])}")
        if exp == 0:
            weaknesses.append("No professional experience detected")
        if achs == 0:
            weaknesses.append("No quantified achievements")
        if certs == 0:
            weaknesses.append("No certifications")
        if not weaknesses:
            weaknesses.append("Minor areas for growth")

        tips = []
        if matched:
            tips.append(f"Highlight experience with {', '.join(matched[:3])}")
        tips.append("Add quantified achievements (e.g., 'Improved by 40%')")
        tips.append(f"Research {company or 'the company'}'s tech stack")
        if missing:
            tips.append(f"Learn {missing[0]}")
        tips.append("Practice STAR method for interviews")

        if missing:
            tw = min(max(len(missing) * 3, 4), 24)
            lp = {
                "totalTimeEstimate": f"{tw - 2}-{tw + 2} weeks",
                "weeklyCommitment": "10-15h",
                "priorityOrder": missing[:6],
                "careerImpact": "Can increase score to 85-95",
                "salaryPotential": f"+${len(missing) * 5}K-${len(missing) * 10}K",
                "jobReadiness": f"Ready in {max(tw // 2, 2)} weeks",
                "milestones": [{"week": 2, "goal": f"Complete {missing[0]} basics", "status": "upcoming"}]
            }
        else:
            lp = {"totalTimeEstimate": "Ready now!", "priorityOrder": [], "jobReadiness": "Immediately ready", "milestones": []}

        gaps = []
        for i, sk in enumerate(missing[:6]):
            if is_soft_skill(sk):
                continue
            sd = get_skill_resources_data(sk)
            gaps.append({
                "skill": sk, "displayName": safe_str(sd.get('displayName'), sk.title()),
                "category": "Technical", "importancePercentage": max(95 - i * 7, 55),
                "priorityLevel": "Critical" if i < 2 else "High" if i < 4 else "Medium",
                "difficulty": "Medium", "learningTime": "4-6 weeks", "demandLevel": "High",
                "salaryImpact": "+$10K", "description": sd.get('description', ''),
                "whyLearn": f"{sk.title()} is required for this role.",
                "topCourses": safe_list(sd.get('freeCoursePlatforms'))[:3],
                "topVideos": safe_list(sd.get('youtubeChannels'))[:3],
                "topProjects": safe_list(sd.get('projectIdeas'))[:3],
            })

        return {
            "overallScore": score, "experienceLevel": level, "fitAnalysis": fit,
            "jobMatchSummary": summary, "hiringRecommendation": rec, "recommendationReason": rr,
            "strengths": strengths, "weaknesses": weaknesses, "interviewTips": tips,
            "overallLearningPlan": lp, "skillGapAnalysis": gaps
        }


# Create analyzer instance
analyzer = ResumeAnalyzer(client if AI_AVAILABLE else None)


# ============================================================================
# API ROUTES
# ============================================================================
@app.route('/api/auth/signup', methods=['POST'])
def signup():
    try:
        d = request.json
        email = d.get('email', '').lower().strip()
        pw = d.get('password', '')
        name = d.get('name', '').strip()
        if not all([email, pw, name]):
            return jsonify({"error": "All fields required"}), 400
        if len(pw) < 6:
            return jsonify({"error": "Password min 6 chars"}), 400
        if DatabaseManager.get_user_by_email(email):
            return jsonify({"error": "Email already registered"}), 400
        u = DatabaseManager.create_user(email, name, pw)
        t = DatabaseManager.create_session(u.id)
        return jsonify({"message": "Signup successful", "user": {"email": u.email, "name": u.name}, "token": t}), 201
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route('/api/auth/login', methods=['POST'])
def login():
    try:
        d = request.json
        email = d.get('email', '').lower().strip()
        pw = d.get('password', '')
        if not all([email, pw]):
            return jsonify({"error": "Email and password required"}), 400
        u = DatabaseManager.get_user_by_email(email)
        if not u or not check_password_hash(u.password, pw):
            return jsonify({"error": "Invalid credentials"}), 401
        t = DatabaseManager.create_session(u.id)
        return jsonify({"message": "Login successful", "user": {"email": u.email, "name": u.name}, "token": t}), 200
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route('/api/auth/logout', methods=['POST'])
def logout():
    DatabaseManager.delete_session(request.headers.get('Authorization', '').replace('Bearer ', ''))
    return jsonify({"message": "Logged out"}), 200


@app.route('/api/auth/me', methods=['GET'])
def get_current_user():
    u = DatabaseManager.get_user_by_token(request.headers.get('Authorization', '').replace('Bearer ', ''))
    if not u:
        return jsonify({"error": "Not authenticated"}), 401
    return jsonify({"user": {"email": u.email, "name": u.name}}), 200


@app.route('/api/analyze', methods=['POST'])
def analyze_resume():
    try:
        u = DatabaseManager.get_user_by_token(request.headers.get('Authorization', '').replace('Bearer ', ''))
        if not u:
            return jsonify({"error": "Please login"}), 401
        if 'resume' not in request.files:
            return jsonify({"error": "No resume uploaded"}), 400
        
        file = request.files['resume']
        jd = request.form.get('jobDescription', '').strip()
        company = request.form.get('companyName', '').strip()
        
        if not jd or len(jd) < 50:
            return jsonify({"error": "Job description too short (min 50 chars)"}), 400

        if file.filename.lower().endswith('.pdf'):
            rt = analyzer.extract_pdf_text(file)
        elif file.filename.lower().endswith('.docx'):
            rt = analyzer.extract_docx_text(file)
        else:
            return jsonify({"error": "Only PDF/DOCX supported"}), 400
        
        if len(rt) < 100:
            return jsonify({"error": "Could not extract resume text."}), 400

        js = analyzer.extract_skills_from_text(jd)
        if not js:
            return jsonify({"error": "No technical skills found in job description."}), 400

        matched, missing = analyzer.match_skills(js, rt)
        rs = analyzer.extract_skills_from_text(rt)
        exp = analyzer.extract_experience(rt)
        score, sb = analyzer.calculate_score(matched, missing, rt, jd, exp, rs)
        analysis = analyzer.generate_analysis(matched, missing, score, exp, company, {"required_skills": js}, sb)

        lr = []
        for sk in missing[:6]:
            if is_soft_skill(sk):
                continue
            sd = get_skill_resources_data(sk)
            lr.append({
                "skillName": sk, "displayName": safe_str(sd.get('displayName'), sk.title()),
                "category": "Technical", "difficulty": "Medium", "learningTime": "4-6 weeks",
                "salaryImpact": "+$10K", "demandLevel": "High", "description": safe_str(sd.get('description')),
                "courses": safe_list(sd.get('freeCoursePlatforms')),
                "youtubeVideos": safe_list(sd.get('youtubeChannels')),
                "projects": safe_list(sd.get('projectIdeas')),
            })

        # Get experience display from details
        exp_details = sb.get('details', {}).get('experience', {})
        exp_display = exp_details.get('totalYearsDisplay', f"{exp}y")

        result = {
            "message": "Analysis complete",
            "companyName": safe_str(company, "Not specified"),
            "overallScore": safe_int(analysis['overallScore']),
            "jobMatchScore": safe_int(sb.get('skillMatchPct', 0)),
            "textSimilarity": safe_int(sb.get('cosineSimilarityPct', 0)),
            "experienceLevel": safe_str(analysis['experienceLevel']),
            "fitAnalysis": safe_str(analysis['fitAnalysis']),
            "jobMatchSummary": safe_str(analysis['jobMatchSummary']),
            "hiringRecommendation": safe_str(analysis['hiringRecommendation']),
            "recommendationReason": safe_str(analysis['recommendationReason']),
            "strengths": safe_list(analysis['strengths']),
            "weaknesses": safe_list(analysis['weaknesses']),
            "interviewTips": safe_list(analysis['interviewTips']),
            "jobRequirements": safe_list(js),
            "requiredSkills": safe_list(js),
            "preferredSkills": [],
            "candidateSkills": safe_list(rs),
            "matchedSkills": safe_list(matched),
            "missingSkills": safe_list(missing),
            "candidateExperience": round(exp, 2),
            "candidateExperienceDisplay": exp_display,
            "requiredExperience": safe_int(sb.get('requiredExperience', 0)),
            "scoreBreakdown": sb,
            "overallLearningPlan": safe_dict(analysis['overallLearningPlan']),
            "skillGapAnalysis": safe_list(analysis['skillGapAnalysis']),
            "learningResources": lr,
            "skillsFound": safe_list(rs),
            "skillsRequired": safe_list(js),
            "skillGaps": safe_list(missing),
            "similarityScore": safe_int(sb.get('cosineSimilarityPct', 0)),
        }

        DatabaseManager.create_analysis(u.id, company, result)
        return jsonify(result), 200

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500


@app.route('/api/health', methods=['GET'])
def health():
    try:
        return jsonify({
            "status": "healthy", "database": "MySQL", "users": User.query.count(),
            "skills": len([s for s in KNOWN_SKILLS if not is_soft_skill(s)]),
            "scoring": "11-metric with accurate month calculation",
            "metrics": ["coreSkillMatch", "skillDepth", "transferableSkills", "cosineSimilarity",
                       "keywordOverlap", "experience", "education", "certifications", "achievements",
                       "resumeQuality", "seniorityMatch"]
        }), 200
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500


@app.route('/api/analyses', methods=['GET'])
def get_user_analyses():
    try:
        token = request.headers.get('Authorization', '').replace('Bearer ', '')
        user = DatabaseManager.get_user_by_token(token)
        if not user:
            return jsonify({"error": "Not authenticated"}), 401

        analyses = Analysis.query.filter_by(user_id=user.id).order_by(Analysis.created_at.desc()).all()
        history = []
        for a in analyses:
            result_data = json.loads(a.result_json) if a.result_json else {}
            history.append({
                "id": a.analysis_id,
                "companyName": a.company_name or "Not specified",
                "overallScore": a.overall_score or 0,
                "jobMatchScore": a.job_match_score or 0,
                "experienceLevel": a.experience_level or "Unknown",
                "createdAt": a.created_at.isoformat() if a.created_at else None,
                "createdAtFormatted": a.created_at.strftime("%B %d, %Y at %I:%M %p") if a.created_at else "Unknown",
                "hiringRecommendation": result_data.get('hiringRecommendation', 'N/A'),
                "matchedSkillsCount": len(result_data.get('matchedSkills', [])),
                "missingSkillsCount": len(result_data.get('missingSkills', [])),
                "totalRequiredSkills": len(result_data.get('jobRequirements', [])),
            })
        return jsonify({"success": True, "count": len(history), "analyses": history}), 200
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500


@app.route('/api/analyses/<analysis_id>', methods=['GET'])
def get_analysis_detail(analysis_id):
    try:
        token = request.headers.get('Authorization', '').replace('Bearer ', '')
        user = DatabaseManager.get_user_by_token(token)
        if not user:
            return jsonify({"error": "Not authenticated"}), 401

        analysis = Analysis.query.filter_by(analysis_id=analysis_id, user_id=user.id).first()
        if not analysis:
            return jsonify({"error": "Analysis not found"}), 404

        result_data = json.loads(analysis.result_json) if analysis.result_json else {}
        result_data['analysisId'] = analysis.analysis_id
        result_data['createdAt'] = analysis.created_at.isoformat() if analysis.created_at else None
        result_data['createdAtFormatted'] = analysis.created_at.strftime("%B %d, %Y at %I:%M %p") if analysis.created_at else "Unknown"
        return jsonify(result_data), 200
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500


@app.route('/api/analyses/<analysis_id>', methods=['DELETE'])
def delete_analysis(analysis_id):
    try:
        token = request.headers.get('Authorization', '').replace('Bearer ', '')
        user = DatabaseManager.get_user_by_token(token)
        if not user:
            return jsonify({"error": "Not authenticated"}), 401

        analysis = Analysis.query.filter_by(analysis_id=analysis_id, user_id=user.id).first()
        if not analysis:
            return jsonify({"error": "Analysis not found"}), 404

        db.session.delete(analysis)
        db.session.commit()
        return jsonify({"success": True, "message": "Analysis deleted"}), 200
    except Exception as e:
        db.session.rollback()
        return jsonify({"error": str(e)}), 500


@app.route('/api/analyses/stats', methods=['GET'])
def get_analysis_stats():
    try:
        token = request.headers.get('Authorization', '').replace('Bearer ', '')
        user = DatabaseManager.get_user_by_token(token)
        if not user:
            return jsonify({"error": "Not authenticated"}), 401

        analyses = Analysis.query.filter_by(user_id=user.id).all()
        if not analyses:
            return jsonify({
                "totalAnalyses": 0, "averageScore": 0, "highestScore": 0, "lowestScore": 0,
                "recentTrend": "neutral", "topCompanies": [],
                "scoreDistribution": {"excellent": 0, "good": 0, "average": 0, "needsWork": 0}
            }), 200

        scores = [a.overall_score for a in analyses if a.overall_score]
        companies = [a.company_name for a in analyses if a.company_name]

        recent_scores = scores[:3] if len(scores) >= 3 else scores
        older_scores = scores[3:6] if len(scores) >= 6 else []
        if older_scores:
            recent_avg = sum(recent_scores) / len(recent_scores)
            older_avg = sum(older_scores) / len(older_scores)
            trend = "improving" if recent_avg > older_avg else "declining" if recent_avg < older_avg else "stable"
        else:
            trend = "neutral"

        distribution = {"excellent": 0, "good": 0, "average": 0, "needsWork": 0}
        for score in scores:
            if score >= 80:
                distribution["excellent"] += 1
            elif score >= 65:
                distribution["good"] += 1
            elif score >= 50:
                distribution["average"] += 1
            else:
                distribution["needsWork"] += 1

        company_counts = Counter(companies)
        top_companies = [{"name": c, "count": n} for c, n in company_counts.most_common(5)]

        return jsonify({
            "totalAnalyses": len(analyses),
            "averageScore": round(sum(scores) / len(scores), 1) if scores else 0,
            "highestScore": max(scores) if scores else 0,
            "lowestScore": min(scores) if scores else 0,
            "recentTrend": trend,
            "topCompanies": top_companies,
            "scoreDistribution": distribution
        }), 200
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500


# ============================================================================
# MAIN ENTRY POINT
# ============================================================================
if __name__ == '__main__':
    print("=" * 60)
    print("RESUME ANALYZER - 11 METRICS WITH ACCURATE CALCULATIONS")
    print("=" * 60)
    print("FIXES APPLIED:")
    print("  ✅ Experience: Accurate month calculation (1y 3m not 2y)")
    print("  ✅ Experience: Section boundary detection")
    print("  ✅ Experience: Returns 0/15 if no experience found")
    print("  ✅ Education: Strict regex with word boundaries")
    print("  ✅ Education: Section boundary detection")
    print("  ✅ Education: Proper field extraction")
    print("=" * 60)
    with app.app_context():
        try:
            db.create_all()
            print("✅ Database ready")
        except Exception as e:
            print(f"⚠️ DB Warning: {e}")
    print(f"🚀 Server: http://localhost:5000")
    print("=" * 60)
    app.run(debug=True, port=5000)